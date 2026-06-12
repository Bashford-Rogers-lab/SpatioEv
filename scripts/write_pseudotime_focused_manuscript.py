"""Write the revised SpatioEv pseudotime manuscript.

This writer follows the current figure logic:

Figure 1: a worked PDAC sample (34434_1).
Figure 2: the pooled four-sample multiplexed imaging analysis.
Figure 3: the Xenium transfer analysis.

Broader toolbox capabilities are framed as helper functions and supplementary
material, keeping the main manuscript focused on morphology/topology-based
epithelial niche pseudotime.
"""

from __future__ import annotations

import json
import sys
from pathlib import Path
from textwrap import dedent

import pandas as pd
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

sys.path.insert(0, str(Path(__file__).resolve().parent))
from write_publication_manuscript import (  # noqa: E402
    AFFILIATION,
    AUTHORS,
    configure_doc,
    fmt_float,
    fmt_int,
    set_cell_margins,
    set_cell_shading,
    set_table_borders,
)


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "manuscript"
FIG_DIR = OUT_DIR / "figures"
SUMMARY_PATH = OUT_DIR / "analysis_summary.json"
DOCX_OUT = OUT_DIR / "SpatioEv_pseudotime_focused_manuscript.docx"
MD_OUT = OUT_DIR / "SpatioEv_pseudotime_focused_manuscript.md"

TITLE = (
    "SpatioEv reconstructs pancreatic epithelial niche evolution from "
    "morphology and topology in spatial omics"
)

KEYWORDS = (
    "spatial pseudotime; multiplexed imaging; pancreatic ductal adenocarcinoma; "
    "PanIN; epithelial niche; tissue topology; desmoplasia; Xenium"
)

MAIN_FIGURES = {
    "Figure 1": "figure_1_34434_pseudotime_workflow.png",
    "Figure 2": "figure_2_pooled_multiplexed_pseudotime.png",
    "Figure 3": "figure_3_xenium_pseudotime_transfer.png",
}

SUPPLEMENTARY_FIGURES = {
    "Supplementary Figure S1": "figure_1_spatioev_workflow.png",
    "Supplementary Figure S2": "figure_2_qc_phenotype_example.png",
    "Supplementary Figure S3": "figure_3_density_spatial_statistics.png",
    "Supplementary Figure S4": "figure_6_ecm_cell_interactions.png",
    "Supplementary Figure S5": "supplementary_xenium_lr_pseudotime_trends.png",
}

REFERENCES = [
    (
        "Palla G, Spitzer H, Klein M, et al. Squidpy: a scalable framework for "
        "spatial omics analysis. Nature Methods 19, 171-178 (2022). "
        "https://www.nature.com/articles/s41592-021-01358-2"
    ),
    (
        "Dries R, Zhu Q, Dong R, et al. Giotto: a toolbox for integrative "
        "analysis and visualization of spatial expression data. Genome Biology "
        "22, 78 (2021). https://genomebiology.biomedcentral.com/articles/10.1186/s13059-021-02286-2"
    ),
    (
        "Marconato L, Palla G, Yamauchi KA, et al. SpatialData: an open and "
        "universal data framework for spatial omics. Nature Methods 22, 58-62 "
        "(2025). https://www.nature.com/articles/s41592-024-02212-x"
    ),
    (
        "Singhal V, Chou N, Lee J, et al. BANKSY unifies cell typing and tissue "
        "domain segmentation for scalable spatial omics data analysis. Nature "
        "Genetics (2024). https://www.nature.com/articles/s41588-024-01664-3"
    ),
    (
        "Goltsev Y, Samusik N, Kennedy-Darling J, et al. Deep profiling of mouse "
        "splenic architecture with CODEX multiplexed imaging. Cell 174, 968-981 "
        "(2018). https://www.cell.com/cell/fulltext/S0092-8674(18)30968-1"
    ),
    (
        "Schurch CM, Bhate SS, Barlow GL, et al. Coordinated cellular "
        "neighborhoods orchestrate antitumoral immunity at the colorectal cancer "
        "invasive front. Cell 182, 1341-1359 (2020). "
        "https://www.cell.com/cell/fulltext/S0092-8674(20)30913-9"
    ),
    (
        "Ligorio M, Sil S, Malagon-Lopez J, et al. Stromal microenvironment "
        "shapes the intratumoral architecture of pancreatic cancer. Cell 178, "
        "160-175 (2019). https://doi.org/10.1016/j.cell.2019.05.012"
    ),
    (
        "Bell ATF, Mitchell JT, Kiemen AL, et al. PanIN and CAF transitions in "
        "pancreatic carcinogenesis revealed with spatial data integration. Cell "
        "Systems 15, 753-769.e5 (2024). https://pmc.ncbi.nlm.nih.gov/articles/PMC11409191"
    ),
    (
        "Braxton AM, Kiemen AL, Grahn MP, et al. 3D genomic mapping reveals "
        "multifocality of human pancreatic precancers. Nature 629, 679-687 "
        "(2024). https://www.nature.com/articles/s41586-024-07359-3"
    ),
    (
        "Agostini A, Piro G, Inzani F, et al. Identification of spatially "
        "resolved markers of malignant transformation in intraductal papillary "
        "mucinous neoplasms. Nature Communications 15, 2764 (2024). "
        "https://www.nature.com/articles/s41467-024-46994-2"
    ),
    (
        "Khaliq AM, Rajamohan M, Masood A, et al. Spatial transcriptomic "
        "analysis of primary and metastatic pancreatic cancers highlights tumor "
        "microenvironmental heterogeneity. Nature Genetics (2024). "
        "https://doi.org/10.1038/s41588-024-01914-4"
    ),
    (
        "Depicting the cellular complexity of pancreatic adenocarcinoma by "
        "imaging mass cytometry: focus on cancer-associated fibroblasts. "
        "Frontiers in Oncology (2024). https://pmc.ncbi.nlm.nih.gov/articles/PMC11578750/"
    ),
    (
        "Multiplexed imaging mass cytometry analysis characterizes the vascular "
        "niche in pancreatic cancer. Oncogenesis (2024). "
        "https://pmc.ncbi.nlm.nih.gov/articles/PMC11250934/"
    ),
]


def p(text: str) -> str:
    return dedent(text).strip().replace("\n", " ")


def load_summary() -> dict:
    if SUMMARY_PATH.exists():
        return json.loads(SUMMARY_PATH.read_text(encoding="utf-8"))
    return {}


def load_stats() -> dict:
    combined = ROOT / "data" / "combined_exp_2_3_4_5"
    pooled = pd.read_pickle(combined / "pooled_niche_result_df.pkl")
    exp2 = pd.read_pickle(combined / "per_sample" / "exp_2_pathology_feature_table.pkl")
    ann = pd.read_csv(ROOT / "data" / "exp_2" / "34434_1_annotation.csv", usecols=["Tier_A", "Tier_B"])
    xenium = pd.read_pickle(ROOT / "data" / "xenium_pancreas_10x" / "pseudotime" / "xenium_pseudotime_result_df.pkl")
    xenium_audit = pd.read_csv(ROOT / "data" / "xenium_pancreas_10x" / "xenium_dataset_audit.csv")
    feature_blocks = pd.read_csv(ROOT / "manuscript" / "analysis_tables" / "supplementary_table_pseudotime_feature_blocks_34434.csv")
    xen_feature_blocks = pd.read_csv(ROOT / "manuscript" / "analysis_tables" / "supplementary_table_xenium_feature_blocks.csv")
    micro = pd.read_csv(
        ROOT
        / "notebooks"
        / "results"
        / "trajectory_microenvironment_interactions"
        / "tables"
        / "multiplexed_microenvironment_trends_contextual.csv"
    )
    interactions = pd.read_csv(
        ROOT
        / "notebooks"
        / "results"
        / "trajectory_microenvironment_interactions"
        / "tables"
        / "multiplexed_epithelial_niche_colocalization_trends.csv"
    )
    xen_micro = pd.read_csv(
        ROOT
        / "notebooks"
        / "results"
        / "trajectory_microenvironment_interactions"
        / "tables"
        / "xenium_microenvironment_trends_sample_centered.csv"
    )

    top_micro = micro.sort_values("spearman_r", key=lambda s: s.abs(), ascending=False).iloc[0]
    top_contact = interactions.sort_values("spearman_r", ascending=False).iloc[0]
    lost_contact = interactions.sort_values("spearman_r").iloc[0]
    top_xen = xen_micro.sort_values("spearman_r", key=lambda s: s.abs(), ascending=False).iloc[0]

    return {
        "pooled": pooled,
        "exp2": exp2,
        "ann": ann,
        "xenium": xenium,
        "xenium_audit": xenium_audit,
        "feature_blocks": feature_blocks,
        "xen_feature_blocks": xen_feature_blocks,
        "micro": micro,
        "interactions": interactions,
        "xen_micro": xen_micro,
        "exp2_cells": len(ann),
        "exp2_niches": len(exp2),
        "pooled_niches": len(pooled),
        "pooled_samples": pooled["sample_id"].nunique(),
        "pooled_branches": pooled["major_branch"].nunique(),
        "xenium_cells": int(xenium_audit["n_cells_matrix"].sum()),
        "xenium_niches": len(xenium),
        "xenium_branches": xenium["major_branch"].nunique(),
        "ductal_cells": int((ann["Tier_A"] == "pancreatic ductal epithelium").sum()),
        "top_micro_label": str(top_micro["label"]),
        "top_micro_r": float(top_micro["spearman_r"]),
        "top_micro_delta": float(top_micro.get("late_minus_early_median", float("nan"))),
        "top_contact_label": str(top_contact["label"]),
        "top_contact_r": float(top_contact["spearman_r"]),
        "lost_contact_label": str(lost_contact["label"]),
        "lost_contact_r": float(lost_contact["spearman_r"]),
        "top_xen_label": str(top_xen["label"]),
        "top_xen_r": float(top_xen["spearman_r"]),
    }


def workflow_table() -> list[tuple[str, str, str]]:
    return [
        (
            "Phenotype scaffold",
            "Tier_A broad cell classes and Tier_B marker-refined states",
            "Establishes epithelial, stromal, immune, endothelial, acinar, and mesenchymal compartments before graph analysis.",
        ),
        (
            "Ductal niche identification",
            "Connected pancreatic ductal epithelial components; minimum five epithelial cells",
            "Treats ducts, glands, duct fragments, and tumor epithelial aggregates as local tissue units.",
        ),
        (
            "Cell graph",
            "30 micrometer radius graph converted by the 0.325 pixel-size scale",
            "Encodes adjacency, epithelial continuity, local mixing, and the physical interface around a ductal niche.",
        ),
        (
            "Niche graph and surroundings",
            "Per-niche topology, geometry, graph-boundary metrics, and five-hop surrounding context",
            "Captures branching, fragmentation, boundary exposure, desmoplasia, immune proximity, and vascular context.",
        ),
        (
            "Pathology modules",
            "Early duct, PanIN-like dysplasia, invasion/desmoplasia, proliferation, and dedifferentiation scores",
            "Turns many raw features into interpretable axes that can be read as tissue biology.",
        ),
        (
            "Trajectory inference",
            "Missingness/variance/correlation filtering, scaling, PCA/UMAP diagnostics, rooted principal tree, tissue back-projection",
            "Orders static niches by similarity to organized ductal, dysplastic, stromal-interface, and invasive-like states.",
        ),
    ]


def helper_table() -> list[tuple[str, str, str]]:
    return [
        (
            "Segmentation QC and phenotype refinement",
            "Area/N:C filters, marker z-scores, SVM-assisted phenotype support, annotation harmonization",
            "Supplementary Figure S2; required before pseudotime because segmentation artifacts distort graph topology.",
        ),
        (
            "Density and point-pattern statistics",
            "Tile densities, phenotype correlations, cross-Ripley, local spatial statistics",
            "Supplementary Figure S3; useful for scale-specific epithelial-stromal and immune colocalization hypotheses.",
        ),
        (
            "Niche boundaries and graph summaries",
            "Connected components, graph neighborhoods, boundary/core contrasts, external edges",
            "Main Figure 1 and supplementary feature tables; these are the central pseudotime building blocks.",
        ),
        (
            "Pseudotime dynamics helpers",
            "Branch-time bins, state cards, transition events, microenvironment trends, interaction trend summaries",
            "Used in Main Figures 1-3 to convert a tree into interpretable biological events.",
        ),
        (
            "ECM-cell analysis",
            "Fiber-object metrics, cell-fiber distances, ECM-proximal cell densities, matrix-cell graph summaries",
            "Supplementary Figure S4; kept outside the main PDAC pseudotime storyline.",
        ),
        (
            "Xenium and spatial transcriptomics extension",
            "Data audit, cell annotation, connected epithelial niches, DAPI/topology/transcript features, BANKSY domain context",
            "Main Figure 3 and Supplementary Figure S5; demonstrates cross-modality transfer with modality-aware limitations.",
        ),
    ]


def content(stats: dict) -> list[dict]:
    return [
        {
            "type": "abstract",
            "title": "Abstract",
            "paragraphs": [
                p(
                    f"""
                    Spatial progression in pancreatic ductal adenocarcinoma
                    (PDAC) is expressed through epithelial architecture, stromal
                    remodeling, immune access, and ductal topology, yet most
                    spatial workflows analyze these signals as separate
                    outputs. We developed SpatioEv to reconstruct epithelial
                    niche pseudotime from static multiplexed images by treating
                    connected ductal epithelial structures and their graph-defined
                    surroundings as the unit of analysis. In a representative
                    PDAC sample, SpatioEv annotated {fmt_int(stats['exp2_cells'])}
                    cells, identified {fmt_int(stats['exp2_niches'])} ductal
                    epithelial niches, extracted morphology, pixel, topology,
                    boundary, and surrounding-context features, converted those
                    features into pathology-inspired modules, and inferred a
                    rooted state trajectory. The same feature grammar generalized
                    to a four-sample multiplexed atlas of {fmt_int(stats['pooled_niches'])}
                    niches, where normal pancreas enriched the trunk and PDAC
                    samples occupied multiple branch programs. Fibroblast
                    abundance was the strongest contextual trend along
                    pseudotime (r = {fmt_float(stats['top_micro_r'])}), and
                    ductal proximity to fibroblasts increased while direct
                    ductal-T-cell excess decreased. A Xenium extension ordered
                    {fmt_int(stats['xenium_niches'])} transcriptomic epithelial
                    niches across {stats['xenium_branches']} branches and
                    reproduced a fibroblast-associated progression signal
                    (r = {fmt_float(stats['top_xen_r'])}). SpatioEv therefore
                    provides a reproducible and interpretable framework for
                    connecting spatial omics measurements to epithelial niche
                    evolution, while broader package functions support QC,
                    density statistics, ECM-cell analysis, and spatial
                    transcriptomics tutorials.
                    """
                )
            ],
        },
        {
            "type": "section",
            "title": "Introduction",
            "paragraphs": [
                p(
                    """
                    PDAC evolves in a tissue ecosystem where epithelial state,
                    fibroblast activation, immune exclusion or engagement,
                    vascular organization, and extracellular matrix remodeling
                    are physically entangled. Recent spatial studies of PanIN,
                    IPMN, primary PDAC, and metastatic PDAC have reinforced this
                    point: precursor lesions are multifocal and heterogeneous,
                    CAF programs emerge near neoplastic epithelium, and spatial
                    transcriptomic maps reveal strong tumor-microenvironmental
                    heterogeneity. These studies motivate a computational
                    question that is not solved by a cell-type map alone: can a
                    static tissue section be ordered into a biologically
                    interpretable sequence of epithelial niche states?
                    """
                ),
                p(
                    """
                    SpatioEv is built around that question. Rather than ordering
                    isolated cells, it first defines a ductal epithelial niche:
                    a connected epithelial structure together with its local
                    graph-defined surroundings. This unit preserves information
                    that pathologists use when reading pancreatic disease,
                    including ductal organization, gland shape, epithelial
                    polarity, nuclear/cell morphology, boundary irregularity,
                    stromal interface exposure, and local immune or fibroblast
                    context. Pseudotime is then used as a state ordering, not as
                    proof of chronological lineage. A high-pseudotime niche is
                    interpreted as being farther from organized duct-like states
                    in morphology/topology/context space, not as a claim that a
                    specific duct was observed transforming over time.
                    """
                ),
                p(
                    """
                    The manuscript is therefore organized around the pseudotime
                    method first. Figure 1 introduces the complete workflow in a
                    representative PDAC sample, 34434. Figure 2 tests the same
                    feature definitions across a pooled four-sample atlas. Figure
                    3 transfers the concept to Xenium, where transcript programs,
                    DAPI morphology, topology, and neighborhood features replace
                    protein-pixel readouts that are unavailable in spatial
                    transcriptomics. The remaining package functions are
                    presented as helper modules and supplementary material.
                    """
                ),
            ],
        },
        {
            "type": "section",
            "title": "Results",
            "paragraphs": [],
        },
        {
            "type": "subsection",
            "title": "A representative PDAC sample establishes the morphology/topology pseudotime workflow",
            "paragraphs": [
                p(
                    f"""
                    We begin with sample 34434 because it contains enough ductal
                    epithelium, stroma, immune cells, and architectural
                    heterogeneity to make every step visible in one field. The
                    Tier_A map separates broad tissue compartments, including
                    {fmt_int(stats['ductal_cells'])} pancreatic ductal epithelial
                    cells, fibroblasts, mesenchymal populations, endothelial
                    cells, T cells, B lineage cells, acinar epithelium, and
                    other compartments. The Tier_B heatmap then shows how these
                    broad classes are refined into marker-defined states such as
                    activated CD4/CD8 T cells and fibroblast subsets defined by
                    aSMA, FAP, Thy1, and PDPN. This two-level phenotyping is not
                    cosmetic: the ductal niche trajectory uses Tier_A to define
                    the epithelial scaffold and surrounding cell classes, while
                    Tier_B provides more granular biological interpretation of
                    immune and stromal context.
                    """
                ),
                p(
                    """
                    Ductal niches are identified as connected components of
                    pancreatic ductal epithelial cells. The connected-component
                    rule is deliberately anatomical: a duct fragment, glandular
                    unit, or tumor epithelial aggregate is treated as a local
                    structure rather than split into unrelated single-cell
                    observations. The examples in Figure 1 show an organized
                    early duct-like niche and a later desmoplastic niche. Around
                    each component, SpatioEv builds a cell graph using a 30
                    micrometer radius and then summarizes the niche itself, its
                    boundary, and graph-defined surrounding hops. This graph
                    makes boundary exposure and stromal interface measurable
                    rather than qualitative.
                    """
                ),
                p(
                    f"""
                    The 34434 feature table contains {fmt_int(int(stats['feature_blocks']['n_features'].sum()))}
                    feature columns across topology, geometry, cell-graph state,
                    graph surroundings, surround composition, cell-state
                    summaries, and PDAC pathology modules. Topology features
                    measure degree, clustering, bridges, connectedness, skeleton
                    leaves, branchpoints, and tortuosity. Geometry features
                    measure compactness, hull shape, orientation, spacing, and
                    boundary irregularity. Pixel and morphology summaries
                    capture nuclear-to-cell ratio, epithelial polarity proxies,
                    texture heterogeneity, entropy, lacunarity, and marker
                    organization. Surrounding-context features quantify
                    fibroblast, immune, endothelial, and mesenchymal composition
                    over graph hops. Together these features make the trajectory
                    interpretable: each axis can be connected to ductal
                    organization, dysplasia-like morphology, desmoplasia,
                    proliferation, or dedifferentiation.
                    """
                ),
                p(
                    """
                    Pathology-inspired modules provide the bridge from hundreds
                    of raw features to biological reading. The early duct anchor
                    is high for compact, connected, polarized duct-like
                    structures. The PanIN-like dysplasia score emphasizes
                    nuclear/cell shape change, polarity loss, and architectural
                    irregularity. The invasion/desmoplasia axis combines
                    epithelial boundary exposure, cross-edges to surrounding
                    cells, fibroblast-rich context, and low circularity.
                    Proliferation and dedifferentiation modules capture Ki67-like
                    growth and high-texture, low-organization states. PCA and
                    UMAP are used as diagnostic views, while the principal tree
                    is rooted in early duct-like niches and projected back to
                    tissue coordinates.
                    """
                ),
                p(
                    """
                    The resulting pseudotime map reconstructs a spatial niche
                    continuum within the static PDAC section. Early states are
                    enriched for organized duct-like structure, intermediate
                    states show PanIN-like and architectural remodeling programs,
                    and later states show desmoplastic and dedifferentiation
                    features. The LOESS trends in Figure 1 are particularly
                    useful because they keep the branch structure visible rather
                    than collapsing the tissue into one linear disease score.
                    Surrounding-context trends show the rise of fibroblast-rich
                    and FAP-associated neighborhoods, while cell-cell interaction
                    dynamics show changing physical contact between ductal cells
                    and fibroblasts, T cells, B lineage cells, and endothelial
                    cells.
                    """
                ),
            ],
        },
        {
            "type": "table",
            "title": "Step-by-step pseudotime workflow used for the 34434 PDAC analysis",
            "headers": ["Step", "Operation", "Biological meaning"],
            "rows": workflow_table(),
        },
        {
            "type": "figure",
            "label": "Figure 1",
            "filename": MAIN_FIGURES["Figure 1"],
            "caption": (
                "Figure 1. A worked PDAC sample introduces morphology/topology "
                "epithelial niche pseudotime. A. Tier_A spatial phenotype map "
                "for sample 34434. B. Tier_B marker-refined phenotypes shown as "
                "a compartment heatmap. C. Early and late ductal niche examples "
                "defined as connected ductal epithelial components, with local "
                "cell graph context. D. Feature families entering cell-graph and "
                "niche-graph summaries. E. PDAC pathology module heatmap over "
                "pseudotime. F-G. PCA and UMAP diagnostics colored by pseudotime. "
                "H. Epithelial niches colored by pseudotime in tissue space. "
                "I. Key pathology score LOESS trends with branch rug. J. "
                "Surrounding-cell programs over pseudotime. K. Ductal cell-cell "
                "contact dynamics within 30 micrometers."
            ),
        },
        {
            "type": "subsection",
            "title": "The pooled multiplexed atlas separates conserved and sample-specific ductal programs",
            "paragraphs": [
                p(
                    f"""
                    We next asked whether the same morphology/topology feature
                    grammar generalizes beyond one PDAC section. The pooled
                    multiplexed atlas contains {fmt_int(stats['pooled_niches'])}
                    ductal epithelial niches from {stats['pooled_samples']}
                    samples: one normal pancreas and three PDAC sections. The
                    pooled UMAP shows both shared structure and sample-dependent
                    occupancy. Normal pancreas is concentrated in the trunk,
                    whereas PDAC niches distribute across multiple branches
                    rather than occupying a single late endpoint. This is
                    biologically important: a PDAC section can contain residual
                    organized duct-like niches, PanIN-like remodeled niches,
                    desmoplastic interface-rich niches, gland-forming tumor
                    architecture, and dedifferentiated-like regions in the same
                    specimen.
                    """
                ),
                p(
                    """
                    The branch-module heatmap turns this pooled structure into a
                    readable biological map. Some branches are early-duct
                    enriched, some are PanIN-like, and others are dominated by
                    invasion/desmoplasia, proliferation, or dedifferentiation.
                    Tissue back-projection confirms that these are not merely
                    embedding artifacts: branch and pseudotime structure map
                    back to real ductal and stromal regions in each sample. This
                    pooled analysis therefore reframes the trajectory as a
                    cross-sample atlas of ductal pathology programs rather than
                    a single-patient timeline.
                    """
                ),
                p(
                    f"""
                    The strongest pooled microenvironment trend was
                    {stats['top_micro_label'].lower()}, which increased with
                    contextual pseudotime (Spearman r =
                    {fmt_float(stats['top_micro_r'])}). The leading positive
                    ductal interaction trend was {stats['top_contact_label'].lower()}
                    (r = {fmt_float(stats['top_contact_r'])}), whereas the
                    strongest negative interaction trend was
                    {stats['lost_contact_label'].lower()} (r =
                    {fmt_float(stats['lost_contact_r'])}). This pattern supports
                    a niche-evolution model in which late contextual states are
                    increasingly fibroblast-associated while some direct
                    epithelial-immune contacts are lost or rearranged in
                    branch-specific ways.
                    """
                ),
            ],
        },
        {
            "type": "figure",
            "label": "Figure 2",
            "filename": MAIN_FIGURES["Figure 2"],
            "caption": (
                "Figure 2. Pooled four-sample analysis separates conserved and "
                "sample-specific ductal niche programs. A. Niche counts across "
                "one normal pancreas and three PDAC samples. B-C. Pooled UMAP "
                "colored by sample and by pseudotime. D. Branch occupancy by "
                "sample. E. Branch-enriched pathology modules. F. Tissue "
                "back-projection of pooled pseudotime for all four samples. "
                "G. Pooled module dynamics across pseudotime. H. Surrounding-cell "
                "trend tests. I. Ductal epithelial interaction trend tests."
            ),
        },
        {
            "type": "subsection",
            "title": "Xenium transfers niche pseudotime to spatial transcriptomics with modality-aware features",
            "paragraphs": [
                p(
                    f"""
                    We then translated the SpatioEv niche concept to Xenium
                    spatial transcriptomics. The Xenium audit includes
                    {fmt_int(stats['xenium_cells'])} cells across one nondiseased
                    pancreas and three PDAC datasets. Because Xenium provides
                    transcripts, cell boundaries, and DAPI morphology rather than
                    multiplexed protein-pixel channels, the feature space is
                    adapted rather than copied. SpatioEv uses connected
                    epithelial components, graph topology, nuclear morphology,
                    transcript marker programs, surrounding cell-type
                    composition, and BANKSY-informed tissue-domain context, while
                    explicitly avoiding claims that require protein polarity
                    markers such as CK19/NaKATPase membrane organization.
                    """
                ),
                p(
                    f"""
                    The pooled Xenium analysis ordered {fmt_int(stats['xenium_niches'])}
                    epithelial niches across {stats['xenium_branches']} major
                    branches. Normal/reference duct-like branches were enriched
                    for nondiseased pancreas, whereas PDAC branches carried
                    combinations of ADM/PanIN-like remodeling, desmoplastic
                    tumor context, immune-inflamed or immune-excluded programs,
                    glandular architecture, and gland-poor/undifferentiated
                    states. The strongest Xenium microenvironment trend was
                    {stats['top_xen_label'].lower()} (r =
                    {fmt_float(stats['top_xen_r'])}), consistent with the
                    fibroblast-associated signal in multiplexed imaging. This
                    cross-modality agreement is encouraging precisely because
                    the two modalities measure different layers of biology.
                    """
                ),
                p(
                    """
                    Figure 3 should therefore be read as a transfer experiment,
                    not a replication with identical features. Multiplexed
                    imaging is stronger for protein state, subcellular polarity,
                    and pixel texture. Xenium is stronger for transcript
                    programs and targeted ligand-receptor potential. SpatioEv
                    provides the common data model that lets both modalities ask
                    the same biological question: how do epithelial structures
                    and their niches vary from organized duct-like states toward
                    remodeled, desmoplastic, immune-altered, or invasive-like
                    tissue states?
                    """
                ),
            ],
        },
        {
            "type": "figure",
            "label": "Figure 3",
            "filename": MAIN_FIGURES["Figure 3"],
            "caption": (
                "Figure 3. Xenium transfers SpatioEv niche pseudotime to spatial "
                "transcriptomics. A. Xenium sample audit showing cell counts and "
                "panel gene counts. B. Tier_A annotation composition by sample. "
                "C. Modality-adapted Xenium feature blocks. D-E. Sample-centered "
                "Xenium niche UMAP colored by sample and pseudotime. F. Spatial "
                "pseudotime maps across Xenium samples. G. Biological program "
                "dynamics across pseudotime. H. Branch biology enrichment. "
                "I. Xenium microenvironment trend tests."
            ),
        },
        {
            "type": "section",
            "title": "Helper functions and supplementary organization",
            "paragraphs": [
                p(
                    """
                    The rest of SpatioEv is best presented as a helper ecosystem
                    around the central pseudotime workflow. These functions are
                    important for users, but they should not compete with the
                    main biological story. The tutorials should show each helper
                    family with a short biological question, expected inputs,
                    generated outputs, and one figure or table demonstrating the
                    answer.
                    """
                )
            ],
        },
        {
            "type": "table",
            "title": "Helper-function families recommended for supplementary/tutorial presentation",
            "headers": ["Helper family", "Functions and outputs", "Where it belongs"],
            "rows": helper_table(),
        },
        {
            "type": "figure",
            "label": "Supplementary Figure S1",
            "filename": SUPPLEMENTARY_FIGURES["Supplementary Figure S1"],
            "caption": (
                "Supplementary Figure S1. Package architecture and reproducible "
                "workflow, including public function groups and tutorial flow."
            ),
        },
        {
            "type": "figure",
            "label": "Supplementary Figure S2",
            "filename": SUPPLEMENTARY_FIGURES["Supplementary Figure S2"],
            "caption": (
                "Supplementary Figure S2. Segmentation QC and phenotype scaffold "
                "for the multiplexed PDAC example."
            ),
        },
        {
            "type": "figure",
            "label": "Supplementary Figure S3",
            "filename": SUPPLEMENTARY_FIGURES["Supplementary Figure S3"],
            "caption": (
                "Supplementary Figure S3. Density, phenotype correlation, and "
                "point-pattern helper functions."
            ),
        },
        {
            "type": "figure",
            "label": "Supplementary Figure S4",
            "filename": SUPPLEMENTARY_FIGURES["Supplementary Figure S4"],
            "caption": (
                "Supplementary Figure S4. ECM-cell helper functions using the "
                "non-PDAC matrix example."
            ),
        },
        {
            "type": "figure",
            "label": "Supplementary Figure S5",
            "filename": SUPPLEMENTARY_FIGURES["Supplementary Figure S5"],
            "caption": (
                "Supplementary Figure S5. Targeted Xenium ligand-receptor "
                "potential trends over sample-centered pseudotime."
            ),
        },
        {
            "type": "section",
            "title": "Discussion",
            "paragraphs": [
                p(
                    """
                    The revised SpatioEv manuscript centers on a stronger claim:
                    static spatial omics can be organized into an interpretable
                    epithelial niche-state trajectory when morphology, topology,
                    boundary exposure, and microenvironment context are measured
                    together. In the worked PDAC example, the trajectory is not a
                    black-box embedding. It is assembled from features that have
                    direct histologic meaning, then summarized as pathology
                    modules that can be visualized in embedding space and in
                    tissue space.
                    """
                ),
                p(
                    """
                    The biological interpretation is coherent with current PDAC
                    spatial literature. Normal or early duct-like organization
                    anchors the trajectory. PanIN-like and architectural
                    complexity programs appear in intermediate states.
                    Fibroblast-rich and FAP-associated surroundings rise with
                    contextual pseudotime, consistent with desmoplastic niche
                    evolution and CAF-associated precursor progression. Direct
                    epithelial-immune contacts do not simply increase or
                    decrease globally; they change by branch, highlighting why a
                    branch-aware spatial trajectory is more informative than a
                    single late-versus-early contrast.
                    """
                ),
                p(
                    """
                    Several limits should remain explicit. Pseudotime orders
                    static niches by state similarity and should not be described
                    as direct temporal observation. PanIN-like, invasive-like,
                    and dedifferentiated-like labels are module interpretations,
                    not diagnostic calls without pathology review. Contextual
                    trajectories include surrounding-cell features by design, so
                    epithelial-intrinsic sensitivity analyses are needed when
                    claiming epithelial-only progression. The WGCNA-like spatial
                    module analysis from earlier drafts is intentionally not a
                    main result until it is implemented, tested, and documented
                    as part of the package API.
                    """
                ),
            ],
        },
        {
            "type": "section",
            "title": "Methods",
            "paragraphs": [
                p(
                    """
                    Multiplexed imaging preprocessing. Single-cell segmentation,
                    marker intensities, phenotype labels, pixel features, and
                    spatial coordinates were aligned by cell label and field of
                    view. Segmentation QC used area and nuclear-to-cell-ratio
                    filters. Tier_A phenotypes defined broad compartments; Tier_B
                    labels captured marker-refined subtypes for immune and
                    stromal interpretation.
                    """
                ),
                p(
                    """
                    Ductal niche and graph construction. Ductal epithelial
                    niches were defined as connected components of pancreatic
                    ductal epithelial cells and retained when at least five
                    epithelial cells were present. Cell graphs were built from
                    spatial coordinates using a 30 micrometer neighborhood radius
                    converted by the 0.325 pixel-size scale. Graph surroundings
                    were summarized over five hops in the representative
                    workflow.
                    """
                ),
                p(
                    """
                    Feature extraction and module scoring. For each niche,
                    SpatioEv summarized morphology, pixel texture, marker state,
                    geometry, graph topology, boundary/core contrasts, external
                    graph degree, cross-edges, and surrounding-cell composition.
                    Signed feature groups were combined into PDAC pathology
                    modules for early duct organization, PanIN-like dysplasia,
                    invasion/desmoplasia, proliferation, and dedifferentiation.
                    """
                ),
                p(
                    """
                    Pseudotime inference and dynamics. Feature matrices were
                    filtered for missingness, imputed, variance-filtered,
                    correlation-filtered, scaled, and inspected by PCA and UMAP.
                    A principal tree was fitted and rooted in organized
                    early-duct-like niches. Pseudotime and branch assignments
                    were projected back onto tissue coordinates. LOESS trends,
                    Spearman correlations, branch-time bins, and epithelial
                    interaction summaries were used to interpret score,
                    microenvironment, and cell-contact dynamics.
                    """
                ),
                p(
                    """
                    Xenium extension. Xenium workflows audited public 10x
                    pancreas datasets, annotated cells with transcript marker
                    programs, constructed connected epithelial niches, and
                    summarized topology, transcript programs, DAPI/nuclear
                    morphology, surrounding composition, and BANKSY domain
                    context. Sample-centered pseudotime and epithelial-intrinsic
                    sensitivity trajectories were used to reduce overinterpretation
                    of sample effects and modality-specific feature availability.
                    """
                ),
            ],
        },
        {
            "type": "section",
            "title": "Data and code availability",
            "paragraphs": [
                p(
                    """
                    The SpatioEv repository contains the package source, tests,
                    tutorials, figure-generation scripts, supplementary feature
                    tables, and manuscript writer. Large local imaging data,
                    H5AD files, Zarr stores, and generated intermediates remain
                    outside Git history under the repository data policy. Public
                    release should provide approved example-data assets or
                    checksums sufficient to rerun the tutorials and regenerate
                    the figures.
                    """
                )
            ],
        },
        {
            "type": "section",
            "title": "References",
            "paragraphs": REFERENCES,
        },
    ]


def add_title_block(doc: Document) -> None:
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(6)
    run = title.add_run(TITLE)
    run.font.name = "Calibri"
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string("0B2545")

    author = doc.add_paragraph()
    author.paragraph_format.space_after = Pt(3)
    run = author.add_run(AUTHORS)
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.bold = True

    aff = doc.add_paragraph()
    aff.paragraph_format.space_after = Pt(10)
    run = aff.add_run(AFFILIATION)
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string("555555")

    keys = doc.add_paragraph()
    keys.paragraph_format.space_after = Pt(10)
    run = keys.add_run(f"Keywords: {KEYWORDS}")
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    run.font.italic = True


def add_key_results_table(doc: Document, stats: dict) -> None:
    rows = [
        ("Representative sample", f"34434: {fmt_int(stats['exp2_cells'])} cells and {fmt_int(stats['exp2_niches'])} ductal epithelial niches"),
        ("Pooled multiplexed atlas", f"{fmt_int(stats['pooled_niches'])} niches from {stats['pooled_samples']} samples and {stats['pooled_branches']} major branches"),
        ("Top contextual trend", f"{stats['top_micro_label']} increases with pseudotime (r = {fmt_float(stats['top_micro_r'])})"),
        ("Top ductal contact trend", f"{stats['top_contact_label']} increases with pseudotime (r = {fmt_float(stats['top_contact_r'])})"),
        ("Strongest lost contact", f"{stats['lost_contact_label']} decreases with pseudotime (r = {fmt_float(stats['lost_contact_r'])})"),
        ("Xenium transfer", f"{fmt_int(stats['xenium_niches'])} niches across {stats['xenium_branches']} branches; {stats['top_xen_label']} r = {fmt_float(stats['top_xen_r'])}"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_borders(table)
    header = table.rows[0].cells
    header[0].text = "Claim"
    header[1].text = "Current example-data support"
    for cell in header:
        set_cell_shading(cell, "F2F4F7")
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    for claim, value in rows:
        cells = table.add_row().cells
        cells[0].text = claim
        cells[1].text = value
        for cell in cells:
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    doc.add_paragraph()


def add_docx_table(doc: Document, block: dict) -> None:
    caption = doc.add_paragraph()
    caption.paragraph_format.keep_with_next = True
    caption.paragraph_format.space_before = Pt(8)
    caption.paragraph_format.space_after = Pt(4)
    run = caption.add_run(block["title"])
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string("1F4D78")

    table = doc.add_table(rows=1, cols=len(block["headers"]))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_borders(table)
    for idx, header in enumerate(block["headers"]):
        cell = table.rows[0].cells[idx]
        cell.text = header
        set_cell_shading(cell, "EAF2F8")
        set_cell_margins(cell, top=90, bottom=90)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(9)
    for row in block["rows"]:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
            set_cell_margins(cells[idx], top=90, bottom=90)
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cells[idx].paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(8.3)
    doc.add_paragraph()


def figure_path(filename: str) -> Path:
    path = FIG_DIR / filename
    if not path.exists():
        raise FileNotFoundError(f"Missing figure: {path}")
    return path


def add_figure_docx(doc: Document, block: dict) -> None:
    path = figure_path(block["filename"])
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run()
    run.add_picture(str(path), width=Inches(6.45))

    cap = doc.add_paragraph()
    cap.paragraph_format.space_after = Pt(10)
    run = cap.add_run(block["caption"])
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor.from_string("333333")


def write_docx(blocks: list[dict], stats: dict) -> None:
    doc = Document()
    configure_doc(doc)
    add_title_block(doc)
    add_key_results_table(doc, stats)
    for block in blocks:
        block_type = block["type"]
        if block_type in {"abstract", "section"}:
            doc.add_heading(block["title"], level=1)
            for text in block["paragraphs"]:
                doc.add_paragraph(text)
        elif block_type == "subsection":
            doc.add_heading(block["title"], level=2)
            for text in block["paragraphs"]:
                doc.add_paragraph(text)
        elif block_type == "table":
            add_docx_table(doc, block)
        elif block_type == "figure":
            add_figure_docx(doc, block)
        else:
            raise ValueError(f"Unknown block type: {block_type}")
    DOCX_OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(DOCX_OUT)


def write_markdown(blocks: list[dict], stats: dict) -> None:
    lines = [
        f"# {TITLE}",
        "",
        AUTHORS,
        "",
        AFFILIATION,
        "",
        f"**Keywords:** {KEYWORDS}",
        "",
        "## Key Results From Current Example Data",
        "",
        "| Claim | Current example-data support |",
        "|---|---|",
        f"| Representative sample | 34434: {fmt_int(stats['exp2_cells'])} cells and {fmt_int(stats['exp2_niches'])} ductal epithelial niches |",
        f"| Pooled multiplexed atlas | {fmt_int(stats['pooled_niches'])} niches from {stats['pooled_samples']} samples and {stats['pooled_branches']} major branches |",
        f"| Top contextual trend | {stats['top_micro_label']} increases with pseudotime (r = {fmt_float(stats['top_micro_r'])}) |",
        f"| Top ductal contact trend | {stats['top_contact_label']} increases with pseudotime (r = {fmt_float(stats['top_contact_r'])}) |",
        f"| Strongest lost contact | {stats['lost_contact_label']} decreases with pseudotime (r = {fmt_float(stats['lost_contact_r'])}) |",
        f"| Xenium transfer | {fmt_int(stats['xenium_niches'])} niches across {stats['xenium_branches']} branches; {stats['top_xen_label']} r = {fmt_float(stats['top_xen_r'])} |",
        "",
    ]
    for block in blocks:
        block_type = block["type"]
        if block_type in {"abstract", "section"}:
            lines.append(f"## {block['title']}")
            lines.append("")
            for text in block["paragraphs"]:
                lines.append(text)
                lines.append("")
        elif block_type == "subsection":
            lines.append(f"### {block['title']}")
            lines.append("")
            for text in block["paragraphs"]:
                lines.append(text)
                lines.append("")
        elif block_type == "table":
            lines.append(f"**{block['title']}**")
            lines.append("")
            lines.append("| " + " | ".join(block["headers"]) + " |")
            lines.append("|" + "|".join(["---"] * len(block["headers"])) + "|")
            for row in block["rows"]:
                lines.append("| " + " | ".join(row) + " |")
            lines.append("")
        elif block_type == "figure":
            rel = figure_path(block["filename"]).relative_to(OUT_DIR).as_posix()
            lines.append(f"![{block['label']}]({rel})")
            lines.append("")
            lines.append(f"*{block['caption']}*")
            lines.append("")
        else:
            raise ValueError(f"Unknown block type: {block_type}")
    MD_OUT.write_text("\n".join(lines).strip() + "\n", encoding="utf-8")


def main() -> None:
    _summary = load_summary()
    stats = load_stats()
    blocks = content(stats)
    write_markdown(blocks, stats)
    write_docx(blocks, stats)
    print(f"Wrote {MD_OUT}")
    print(f"Wrote {DOCX_OUT}")


if __name__ == "__main__":
    main()
