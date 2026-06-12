"""Write the current-package SpatioEv manuscript draft as DOCX and Markdown."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "manuscript"
DOCX_OUT = OUT_DIR / "SpatioEv_current_package_manuscript.docx"
MD_OUT = OUT_DIR / "SpatioEv_current_package_manuscript.md"


TITLE = (
    "SpatioEv: a reproducible Python toolbox for spatial evolution analysis "
    "of multiplexed imaging and spatial transcriptomics data"
)

AUTHORS = (
    "Shihong Wu, Sakina Amin, Liezel Tamon, Rachael Bashford-Rogers, "
    "and colleagues"
)


SECTIONS = [
    (
        "Abstract",
        [
            (
                "Spatial biology increasingly requires analyses that preserve "
                "single-cell identity, tissue architecture, extracellular matrix "
                "organization, and disease-associated transitions in the same "
                "framework. We present SpatioEv, a modular Python package for "
                "spatial feature extraction and spatial evolution analysis in "
                "multiplexed imaging data, with extensions to spatial "
                "transcriptomics. The current package provides segmentation "
                "quality control, marker and morphology feature engineering, "
                "semi-supervised phenotype refinement, density and point-pattern "
                "statistics, spatial autocorrelation, cell-cell and cell-ECM "
                "interaction analysis, niche boundary and cell-graph summaries, "
                "pixel-level morphology extraction, and pseudotime-linked "
                "microenvironment dynamics. We reorganized the codebase into a "
                "GitHub-ready package with lazy optional dependencies, explicit "
                "metadata, reproducible tests, a data policy for large imaging "
                "assets, and tutorial notebooks that run on either local example "
                "data or synthetic fallbacks. Using previous disease-focused "
                "applications as biological reference cases, SpatioEv enables "
                "systematic interrogation of quality, phenotype confidence, "
                "multi-scale tissue organization, ECM-cell coupling, and "
                "niche-level progression. The WGCNA-like spatial feature module "
                "described in earlier drafts is not part of this release and is "
                "therefore reserved for future integration."
            )
        ],
    ),
    (
        "Introduction",
        [
            (
                "Cells function within spatially structured tissues, where "
                "cellular identity, morphology, local neighbors, extracellular "
                "matrix composition, and tissue boundaries jointly shape "
                "biological behavior. Dissociation-based single-cell methods "
                "provide molecular resolution but remove the spatial "
                "relationships needed to study immune infiltration, stromal "
                "organization, tumor invasion, epithelial polarity, and matrix "
                "remodeling. Highly multiplexed imaging and spatial "
                "transcriptomics retain this organization and can profile large "
                "numbers of cells in situ, but they require robust computational "
                "methods that can handle segmentation errors, annotation "
                "uncertainty, heterogeneous image quality, and multi-scale "
                "spatial structure."
            ),
            (
                "Existing spatial analysis workflows often emphasize cell type "
                "composition, pairwise proximity, or region discovery. These are "
                "important but incomplete views of tissue organization. Many "
                "questions require simultaneous analysis of cell morphology, "
                "nuclear and membrane texture, phenotype probability, density "
                "variation, tissue boundaries, ECM fibers, graph topology, and "
                "pseudotime-like transitions. A reusable toolbox should also "
                "separate core methods from heavy viewer or single-cell "
                "dependencies so users can install and test the package in a "
                "plain Python environment before enabling optional imaging and "
                "scverse integrations."
            ),
            (
                "SpatioEv addresses this need as an extensible Python package "
                "organized around reusable modules rather than one-off analysis "
                "notebooks. The current release supports multiplexed imaging as "
                "the primary input modality and includes additional utilities "
                "for Xenium-derived spatial transcriptomics analyses. The "
                "package is designed to complement AnnData-based workflows and "
                "to expose interpretable intermediate tables that can be audited "
                "and reused in manuscripts, quality-control reports, and "
                "downstream biological models."
            ),
        ],
    ),
    (
        "Results",
        [],
    ),
    (
        "A reusable and reproducible package release",
        [
            (
                "We reorganized SpatioEv into an installable Python package with "
                "a clear public API, packaging metadata, optional dependency "
                "groups, automated tests, and cleaned tutorial notebooks. The "
                "top-level import is now lightweight: importing spatioev no "
                "longer loads Scanpy, scimap, Napari, or other heavy optional "
                "tools. Instead, heavy dependencies are imported only when a "
                "dependent function is called. This makes the package easier to "
                "test in continuous integration and safer to use in constrained "
                "analysis environments."
            ),
            (
                "The repository is now structured for GitHub upload. Source "
                "code lives under spatioev/, smoke tests under tests/, cleaned "
                "tutorials under tutorials/, and release notes under docs/. "
                "Large local imaging assets and generated results are excluded "
                "by a data policy and .gitignore rules. The test suite includes "
                "deterministic toy-data tests for core APIs and a local smoke "
                "test that uses data/exp_2/34434_1_adata.h5ad when available, "
                "while skipping cleanly in lightweight checkouts."
            ),
        ],
    ),
    (
        "Quality control and phenotype feature engineering",
        [
            (
                "SpatioEv begins with segmentation quality control because "
                "segmentation errors can propagate into every downstream "
                "spatial statistic. The qc module computes cell area in physical "
                "units, classifies debris-like and merged-cell-like objects, "
                "flags abnormal nuclear-to-cell ratios, and summarizes removal "
                "rates globally or by image. These functions operate directly on "
                "AnnData observation metadata and return auditable columns rather "
                "than hidden state."
            ),
            (
                "For phenotype analysis, SpatioEv combines marker intensities "
                "with morphology descriptors. The preprocessing module can copy "
                "selected marker values from the expression matrix into obs and "
                "z-score numeric features. The ml module constructs marker and "
                "morphology feature matrices, supports the historical spelling "
                "variant fractual_dimension present in some Pixie-derived data, "
                "and trains radial-basis SVM classifiers to produce both class "
                "labels and phenotype probabilities. These probability profiles "
                "support annotation refinement and transition-state inspection "
                "without discarding the original annotations."
            ),
        ],
    ),
    (
        "Multi-scale spatial density and interaction statistics",
        [
            (
                "The spatial module quantifies tissue organization at several "
                "scales. Tile-based density functions measure object count and "
                "pixel-area occupancy across an image grid, while k-nearest "
                "neighbor and fixed-radius density functions provide per-cell "
                "local density estimates. Phenotype-specific density tables can "
                "be correlated across spatial tiles to identify co-enriched or "
                "mutually exclusive cellular neighborhoods."
            ),
            (
                "SpatioEv also implements point-pattern and autocorrelation "
                "statistics that help convert spatial impressions into "
                "quantitative evidence. Ripley and cross-Ripley functions "
                "measure clustering and phenotype attraction over user-defined "
                "radii. Moran and cross-Moran functions measure spatial "
                "autocorrelation of continuous features and coupling between "
                "features. Local versions add per-cell hotspot or co-localization "
                "scores back to AnnData, enabling downstream visualization and "
                "phenotype-specific summaries."
            ),
        ],
    ),
    (
        "Niche boundaries and graph-based tissue representations",
        [
            (
                "Many biologically important events occur at interfaces between "
                "cellular regions, including tumor-stroma boundaries, immune "
                "aggregates, ducts, and peri-epithelial niches. SpatioEv provides "
                "functions for clustering spatial components, constructing "
                "niche boundaries, buffering regions, assigning cells to niche "
                "regions, and summarizing niche composition. The current package "
                "also builds cell graphs from spatial coordinates and selected "
                "cell features, creating adjacency and distance matrices that "
                "can be used to summarize local topology, feature organization, "
                "boundary-core differences, and surrounding tissue context."
            ),
            (
                "These graph features extend earlier region-of-interest and "
                "boundary analyses by making niche representations reproducible "
                "and comparable across images. In previous applications, this "
                "type of analysis supported automated tumor nest selection, "
                "boundary characterization, and recurrent cellular neighborhood "
                "summaries. In the current package, the reusable functions are "
                "available as explicit API calls and are covered by tutorial "
                "workflows rather than being embedded only in development "
                "notebooks."
            ),
        ],
    ),
    (
        "Integrated ECM-cell spatial analysis",
        [
            (
                "The ECM analysis modules link fiber-level summaries to "
                "single-cell phenotypes. SpatioEv includes utilities to build "
                "cell-fiber links, map nearest fibers to cells, quantify fiber "
                "density around cells, compute cross-Ripley and cross-Moran "
                "statistics between cells and ECM features, map cell-derived "
                "signals onto fibers, and estimate fiber-cell alignment. These "
                "functions allow extracellular matrix structure to be analyzed "
                "as part of the tissue ecosystem rather than as a separate image "
                "layer."
            ),
            (
                "Earlier RA and OA analyses used this logic to study how fiber "
                "types and fiber morphology relate to immune and stromal "
                "cellular neighborhoods. The current package preserves the core "
                "ECM-cell operations in spatioev.spatial while keeping the "
                "high-dimensional WGCNA-like module analysis outside the public "
                "release until it can be implemented, tested, and documented as "
                "a first-class package component."
            ),
        ],
    ),
    (
        "Spatial trajectory and spatial transcriptomics extensions",
        [
            (
                "SpatioEv supports niche-level progression analysis by combining "
                "morphology, neighborhood composition, and pseudotime assignments. "
                "The pseudotime_dynamics module bins continuous pseudotime values, "
                "computes source-centered interaction metrics for target "
                "phenotypes along pseudotime, and summarizes how local "
                "microenvironmental relationships change across trajectory "
                "states. This makes it possible to ask whether epithelial, "
                "immune, stromal, or ECM-associated features change gradually, "
                "branch specifically, or abruptly across disease-associated "
                "progression."
            ),
            (
                "The current repository also includes Xenium-oriented scripts and "
                "notebooks for data audit, cell annotation, epithelial niche "
                "feature extraction, pooled pseudotime, BANKSY integration, and "
                "SpatialCellChat integration. These workflows expand SpatioEv "
                "from multiplexed protein imaging toward spatial transcriptomics "
                "while retaining the same design principle: each step should "
                "produce reusable tables and figures that can be audited, tested, "
                "and connected to downstream biological interpretation."
            ),
        ],
    ),
    (
        "Discussion",
        [
            (
                "The current SpatioEv package converts a development workspace "
                "into a reusable spatial-analysis toolbox. Its main advance is "
                "not a single statistic, but the integration of quality control, "
                "phenotype confidence, spatial density, point-pattern statistics, "
                "ECM-cell coupling, niche topology, and pseudotime-linked "
                "microenvironment dynamics into a modular API. This structure "
                "allows users to inspect intermediate results and combine modules "
                "according to the biological system rather than follow a single "
                "rigid pipeline."
            ),
            (
                "The package is especially suited to whole-slide or large-field "
                "multiplexed imaging experiments, where millions of cells and "
                "large image-derived feature tables make manual inspection and "
                "ad hoc analysis fragile. Lazy optional imports, data-aware tests, "
                "and runnable tutorials also make the toolbox more reproducible "
                "for collaborators who may not have the original 129 GB local "
                "workspace. This is important for translational spatial biology, "
                "where reproducibility depends as much on data stewardship and "
                "installation clarity as on statistical method choice."
            ),
            (
                "Several limitations remain. Full biological validation still "
                "depends on accurate segmentation, marker panel design, image "
                "registration, and sample metadata. Some high-level workflows, "
                "including the WGCNA-like spatial feature module analysis from "
                "earlier drafts, remain outside the current package API. In "
                "addition, some optional workflows require Scanpy, scimap, "
                "Napari, SpatialData, Squidpy, ElPiGraph, BANKSY, or R-based "
                "tools that should be documented and tested separately. Future "
                "work should formalize these extensions, add continuous "
                "integration with public demo data, and provide versioned example "
                "outputs for manuscript figures."
            ),
            (
                "By making spatial feature extraction reproducible and modular, "
                "SpatioEv provides a foundation for studying how tissue "
                "architecture evolves across inflammation, cancer progression, "
                "stromal remodeling, and spatial transcriptomic states. The "
                "current release is therefore best viewed as a stable toolbox "
                "core: ready for GitHub upload, tutorial-driven adoption, and "
                "incremental extension as additional workflows mature."
            ),
        ],
    ),
    (
        "Materials and Methods",
        [],
    ),
    (
        "Package implementation",
        [
            (
                "SpatioEv is implemented as a Python package named spatioev. The "
                "package uses pyproject.toml metadata with setuptools, requires "
                "Python 3.10 or newer, and defines optional extras for developer "
                "tools, Scanpy-based workflows, scimap/Napari viewers, and "
                "SpatialData/Squidpy workflows. Core dependencies include AnnData, "
                "NumPy, pandas, SciPy, scikit-learn, scikit-image, matplotlib, "
                "seaborn, NetworkX, statsmodels, tqdm, and Shapely. Optional "
                "dependencies are imported lazily at function call time."
            )
        ],
    ),
    (
        "Input data model",
        [
            (
                "Most functions operate on AnnData objects containing single-cell "
                "measurements in X, feature names in var_names, and per-cell "
                "metadata in obs. Standard coordinate columns are X_centroid and "
                "Y_centroid, and image membership is stored in imageid. Many "
                "functions also accept explicit column-name arguments so users "
                "can adapt the methods to other naming conventions."
            )
        ],
    ),
    (
        "Segmentation quality control",
        [
            (
                "Cell area is converted into square microns using a user-defined "
                "pixel size. Cells are classified as debris-like, merged-cell-like, "
                "or normal based on minimum and maximum area thresholds. Cells "
                "with nuclear-to-cell ratio above a configurable threshold are "
                "flagged as abnormal. Summary tables can be produced globally or "
                "stratified by image or sample."
            )
        ],
    ),
    (
        "Feature engineering and SVM phenotyping",
        [
            (
                "Marker features are extracted from selected AnnData variables "
                "and standardized. Morphology features include area, convex area, "
                "perimeter, axis lengths, Feret diameter, equivalent diameter, "
                "concavity count, centroid difference, eccentricity, solidity, "
                "axis ratios, circularity, boundary irregularity, fractal "
                "dimension, and nuclear-to-cell ratio. The combined matrix is "
                "used to train an RBF-kernel SVM with class balancing and "
                "probability estimates. Predicted labels and class probabilities "
                "are written back to obs."
            )
        ],
    ),
    (
        "Spatial density and spatial statistics",
        [
            (
                "Tile-based density partitions each image into square bins and "
                "computes object density and pixel-area density. Local density is "
                "computed either as the inverse mean k-nearest-neighbor distance "
                "or as neighbor count per fixed-radius area. Phenotype interaction "
                "density counts target-phenotype cells around source-phenotype "
                "cells within a radius."
            ),
            (
                "Ripley functions use BallTree neighbor queries and convex-hull "
                "window areas to compute K, expected K, L, and L-minus-r values. "
                "Cross-Ripley functions compare source and target cell sets and "
                "add source-centered neighborhood summaries. Moran functions use "
                "k-nearest-neighbor graphs to estimate global and local spatial "
                "autocorrelation for individual features and paired features."
            ),
        ],
    ),
    (
        "Niche, graph, ECM, and pseudotime workflows",
        [
            (
                "Niche boundary functions identify spatial components, construct "
                "boundaries, buffer regions, assign cells to regions, and "
                "summarize niche composition. Cell graph functions build spatial "
                "graphs within images from coordinates and feature vectors, then "
                "extract induced subgraphs and niche graph summaries. ECM "
                "functions quantify cell-to-fiber distances, fiber density near "
                "cells, cross-statistics between cell phenotypes and fibers, "
                "fiber alignment, and ECM niche graphs. Pseudotime functions bin "
                "trajectory values and summarize source-centered interaction "
                "dynamics across bins."
            )
        ],
    ),
    (
        "Testing and tutorial verification",
        [
            (
                "The package includes pytest tests for imports, QC, feature "
                "engineering, density, interactions, Ripley and Moran statistics, "
                "pseudotime dynamics, and the local exp_2 example dataset. Four "
                "tutorial notebooks were generated and executed successfully in a "
                "temporary directory: setup/data audit, QC and phenotype features, "
                "spatial density and statistics, and niche graph/pseudotime/"
                "Xenium extensions."
            )
        ],
    ),
    (
        "Code and data availability",
        [
            (
                "The package is prepared for release at "
                "https://github.com/Bashford-Rogers-lab/SpatioEv. The local "
                "working directory contains large raw and derived imaging data "
                "that are excluded from GitHub upload. The repository includes a "
                "data policy describing how local H5AD, OME-TIFF, Zarr, pickle, "
                "and results files should be handled. Public example data should "
                "be distributed separately with checksums or archive accessions."
            )
        ],
    ),
    (
        "Acknowledgements",
        [
            (
                "S.W. is supported by Cancer Research UK. S.A. is funded by the "
                "St John Clarendon Fund scholarship. E.C. is supported by NIHR "
                "BRC-Oxford. We acknowledge the Oxford Centre for Histopathology "
                "Research, Oxford Radcliffe Biobank, the University of Oxford, "
                "Oxford CRUK Cancer Centre, NIHR Oxford Biomedical Research "
                "Centre, and NIHR CRN Thames Valley network. Author affiliations, "
                "ethics statements, and funding details should be checked against "
                "the final submission files."
            )
        ],
    ),
    (
        "Author contributions",
        [
            (
                "S.W., S.A., K.M., and R.B.-R. conceived and designed the "
                "analysis. S.W., S.A., and R.B.-R. performed analyses and package "
                "development with contributions from collaborators. Dataset "
                "collection, sample preparation, imaging, biological "
                "interpretation, manuscript writing, and supervision should be "
                "finalized using the project authorship spreadsheet before "
                "submission."
            )
        ],
    ),
    (
        "Declaration of interests",
        [
            (
                "R.B.-R. is a co-founder of Alchemab Therapeutics Ltd and has "
                "consulting relationships listed in the previous manuscript. E.C. "
                "has consulting relationships listed in the previous manuscript. "
                "All declarations should be reviewed and updated before "
                "submission."
            )
        ],
    ),
    (
        "Ethics approval",
        [
            (
                "Informed consent was obtained for patient samples where "
                "applicable. The PDAC Surrey cohort was retrieved with ethical "
                "approval under IRAS project 277406, ethical reference "
                "20/SW/0105. HNSCC tissue was retrieved with ethical approval "
                "under IRAS project 262470, ethical reference 19/SC/0173. RA and "
                "OA ethics details should be inserted from the final clinical "
                "metadata files."
            )
        ],
    ),
    (
        "Figure legends",
        [
            (
                "Figure 1. SpatioEv package overview and quality-control "
                "workflow. The optimized package exposes reusable modules for "
                "segmentation QC, phenotype feature engineering, spatial density, "
                "point-pattern statistics, niche graphs, ECM-cell analysis, and "
                "pseudotime-linked microenvironment dynamics."
            ),
            (
                "Figure 2. Phenotype refinement with marker and morphology "
                "features. SVM-derived probability profiles support annotation "
                "quality control, reassignment of low-confidence labels, and "
                "identification of transitional cell states."
            ),
            (
                "Figure 3. Multi-scale density and spatial statistics. Tile, "
                "radius, kNN, Ripley, and Moran analyses quantify cellular "
                "density, phenotype co-enrichment, clustering, and feature "
                "autocorrelation across tissue scales."
            ),
            (
                "Figure 4. Niche boundaries and graph summaries. Spatial "
                "components are converted into boundaries, buffered regions, "
                "cell assignments, and graph-derived summaries of topology, "
                "composition, and surrounding context."
            ),
            (
                "Figure 5. ECM-cell spatial analysis. Fiber-level properties are "
                "linked to cell phenotypes to measure local fiber density, "
                "cell-fiber proximity, alignment, and feature coupling."
            ),
            (
                "Figure 6. Spatial trajectory and Xenium extensions. Niche-level "
                "features and local microenvironment metrics can be summarized "
                "over pseudotime, with related workflows extending the package "
                "toward Xenium annotation, niche features, BANKSY, and "
                "SpatialCellChat integration."
            ),
            (
                "Previous Figure 7 WGCNA-like spatial feature modules are not "
                "included in the current package release and should remain out "
                "of the main current-package manuscript until implemented and "
                "tested as package functions."
            ),
        ],
    ),
    (
        "References",
        [
            (
                "References should be rebuilt from the final target journal style. "
                "At minimum, include citations for AnnData, Scanpy, scikit-learn, "
                "scikit-image, NetworkX, Shapely, SpatialData, Squidpy, scimap, "
                "Napari, ElPiGraph, BANKSY, SpatialCellChat, Ripley statistics, "
                "Moran statistics, and the biological literature cited in the "
                "previous disease-specific results."
            )
        ],
    ),
]


MODULE_TABLE = [
    ["Module", "Current package role"],
    ["spatioev.qc", "Segmentation QC, area and nuclear-to-cell ratio summaries"],
    ["spatioev.preprocessing", "Marker-to-obs transfer and z-score feature creation"],
    ["spatioev.ml", "Marker/morphology feature matrices and SVM phenotype models"],
    ["spatioev.spatial.general_density", "Tile, phenotype, KDE, and correlation density summaries"],
    ["spatioev.spatial.spatial_stats", "Ripley, Moran, local hotspot, and cross-feature statistics"],
    ["spatioev.spatial.spatial_niche_boundaries", "Component clustering, boundaries, buffers, and niche membership"],
    ["spatioev.spatial.spatial_cell_graph", "Cell graph construction and niche subgraph extraction"],
    ["spatioev.spatial.spatial_ecm_*", "Cell-fiber links, ECM neighborhoods, and ECM graph summaries"],
    ["spatioev.spatial.pseudotime_dynamics", "Pseudotime bins and source-centered interaction dynamics"],
    ["scripts and tutorials", "Xenium, BANKSY, SpatialCellChat, and reproducible worked examples"],
]


def set_cell_width(cell, width_dxa: int) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "4")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), "D9DEE7")
        borders.append(tag)


def set_cell_fill(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_style_font(style, size_pt: float, color: str | None = None, bold: bool = False) -> None:
    style.font.name = "Calibri"
    style.font.size = Pt(size_pt)
    style.font.bold = bold
    if color:
        style.font.color.rgb = RGBColor.from_string(color)


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.start_type = WD_SECTION_START.NEW_PAGE

    normal = doc.styles["Normal"]
    set_style_font(normal, 11)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333

    for style_name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = doc.styles[style_name]
        set_style_font(style, size, color=color, bold=True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.208


def add_title(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(TITLE)
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor.from_string("0B2545")

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(AUTHORS)
    run.italic = True
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string("555555")


def add_module_table(doc: Document) -> None:
    doc.add_heading("Current package module map", level=2)
    table = doc.add_table(rows=len(MODULE_TABLE), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_borders(table)
    widths = [3120, 6240]
    for r_idx, row in enumerate(MODULE_TABLE):
        for c_idx, value in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_width(cell, widths[c_idx])
            if r_idx == 0:
                set_cell_fill(cell, "F4F6F9")
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(value)
            run.font.name = "Calibri"
            run.font.size = Pt(9.5)
            if r_idx == 0:
                run.bold = True


def build_markdown() -> str:
    lines = [f"# {TITLE}", "", AUTHORS, ""]
    for heading, paragraphs in SECTIONS:
        level = "##" if heading in {"Abstract", "Introduction", "Results", "Discussion", "Materials and Methods", "Code and data availability", "Acknowledgements", "Author contributions", "Declaration of interests", "Ethics approval", "Figure legends", "References"} else "###"
        lines.append(f"{level} {heading}")
        lines.append("")
        for para in paragraphs:
            lines.append(para)
            lines.append("")
        if heading == "A reusable and reproducible package release":
            lines.append("### Current package module map")
            lines.append("")
            header = MODULE_TABLE[0]
            lines.append(f"| {header[0]} | {header[1]} |")
            lines.append("|---|---|")
            for row in MODULE_TABLE[1:]:
                lines.append(f"| {row[0]} | {row[1]} |")
            lines.append("")
    return "\n".join(lines).strip() + "\n"


def build_docx() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_styles(doc)
    add_title(doc)

    for heading, paragraphs in SECTIONS:
        if heading in {"Abstract", "Introduction", "Results", "Discussion", "Materials and Methods", "Code and data availability", "Acknowledgements", "Author contributions", "Declaration of interests", "Ethics approval", "Figure legends", "References"}:
            doc.add_heading(heading, level=1)
        else:
            doc.add_heading(heading, level=2)

        for para in paragraphs:
            doc.add_paragraph(para)

        if heading == "A reusable and reproducible package release":
            add_module_table(doc)

    doc.save(DOCX_OUT)
    MD_OUT.write_text(build_markdown(), encoding="utf-8")
    print(DOCX_OUT)
    print(MD_OUT)


if __name__ == "__main__":
    build_docx()
