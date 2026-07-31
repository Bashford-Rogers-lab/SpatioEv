"""Write a publication-style SpatioEv manuscript with generated figures.

The manuscript is intentionally built from the current package state and the
local example-analysis outputs, rather than from static prose. Re-run
``scripts/generate_manuscript_figures.py`` before this script when the example
datasets or figure tables change.
"""

from __future__ import annotations

import json
from pathlib import Path
from textwrap import dedent

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "manuscript"
FIG_DIR = OUT_DIR / "figures"
SUMMARY_PATH = OUT_DIR / "analysis_summary.json"

DOCX_OUT = OUT_DIR / "SpatioEv_publication_manuscript.docx"
MD_OUT = OUT_DIR / "SpatioEv_publication_manuscript.md"

TITLE = (
    "SpatioEv links tissue architecture, extracellular matrix organization, "
    "and spatial evolution in multiplexed imaging and spatial transcriptomics"
)

AUTHORS = (
    "Shihong Wu, Sakina Amin, Liezel Tamon, Rachael Bashford-Rogers, "
    "and colleagues"
)

AFFILIATION = "Bashford-Rogers laboratory, University of Oxford"

KEYWORDS = (
    "spatial biology; multiplexed imaging; pancreatic cancer; extracellular "
    "matrix; Xenium; spatial transcriptomics; tissue microenvironment"
)


REFERENCES = [
    (
        "Palla G, Spitzer H, Klein M, et al. Squidpy: a scalable framework for "
        "spatial omics analysis. Nature Methods 19, 171-178 (2022). "
        "https://www.nature.com/articles/s41592-021-01358-2"
    ),
    (
        "Dries R, Zhu Q, Dong R, et al. Giotto: a toolbox for integrative "
        "analysis and visualization of spatial expression data. Genome Biology "
        "22, 78 (2021). "
        "https://genomebiology.biomedcentral.com/articles/10.1186/s13059-021-02286-2"
    ),
    (
        "Singhal V, Chou N, Lee J, et al. BANKSY unifies cell typing and tissue "
        "domain segmentation for scalable spatial omics data analysis. Nature "
        "Genetics (2024). https://www.nature.com/articles/s41588-024-01664-3"
    ),
    (
        "Marconato L, Palla G, Yamauchi KA, et al. SpatialData: an open and "
        "universal data framework for spatial omics. Nature Methods 22, 58-62 "
        "(2025). https://www.nature.com/articles/s41592-024-02212-x"
    ),
    (
        "Goltsev Y, Samusik N, Kennedy-Darling J, et al. Deep profiling of mouse "
        "splenic architecture with CODEX multiplexed imaging. Cell 174, 968-981 "
        "(2018). https://www.cell.com/cell/fulltext/S0092-8674(18)30968-1"
    ),
    (
        "Schurch CM, Bhate SS, Barlow GL, et al. Coordinated cellular "
        "neighborhoods orchestrate antitumoral immunity at the colorectal cancer "
        "invasive front. Cell 182, 1341-1359 (2020). "
        "https://www.cell.com/cell/fulltext/S0092-8674(20)30913-9"
    ),
    (
        "Ligorio M, Sil S, Malagon-Lopez J, et al. Stromal microenvironment "
        "shapes the intratumoral architecture of pancreatic cancer. Cell 178, "
        "160-175 (2019). https://doi.org/10.1016/j.cell.2019.05.012"
    ),
    (
        "PanIN and CAF transitions in pancreatic carcinogenesis revealed with "
        "spatial data integration. PMC article. "
        "https://pmc.ncbi.nlm.nih.gov/articles/PMC11409191"
    ),
    (
        "Liot S, Balas J, Aubert A, et al. Stroma involvement in pancreatic "
        "ductal adenocarcinoma: an overview focusing on extracellular matrix "
        "proteins. Frontiers in Immunology 12, 612271 (2021). "
        "https://pubmed.ncbi.nlm.nih.gov/33889150/"
    ),
    (
        "Matsubara T, Spycher MA, Ruttner JR, Fehr K. The localization of "
        "fibronectin in rheumatoid arthritis synovium by light and electron "
        "microscopic immunohistochemistry. Rheumatology International 3, 153-159 "
        "(1983). https://pubmed.ncbi.nlm.nih.gov/6366989/"
    ),
]


def load_summary() -> dict:
    if not SUMMARY_PATH.exists():
        raise FileNotFoundError(
            f"{SUMMARY_PATH} is missing. Run scripts/generate_manuscript_figures.py first."
        )
    return json.loads(SUMMARY_PATH.read_text(encoding="utf-8"))


def fmt_int(value: int | float) -> str:
    return f"{int(value):,}"


def fmt_float(value: float, digits: int = 3) -> str:
    return f"{float(value):.{digits}f}"


def fmt_pct(value: float, digits: int = 1) -> str:
    return f"{float(value):.{digits}f}%"


def figure_path(number: int) -> Path:
    names = {
        1: "figure_1_spatioev_workflow.png",
        2: "figure_2_qc_phenotype_example.png",
        3: "figure_3_density_spatial_statistics.png",
        4: "figure_4_pseudotime_niche_biology.png",
        5: "figure_5_microenvironment_dynamics.png",
        6: "figure_6_ecm_cell_interactions.png",
        7: "figure_7_xenium_extension.png",
    }
    path = FIG_DIR / names[number]
    if not path.exists():
        raise FileNotFoundError(f"Missing figure: {path}")
    return path


def build_content(summary: dict) -> list[dict]:
    f1 = summary["figure_1"]
    f2 = summary["figure_2"]
    f3 = summary["figure_3"]
    f4 = summary["figure_4"]
    f5 = summary["figure_5"]
    f6 = summary["figure_6"]
    f7 = summary["figure_7"]

    top_density = f3["top_positive_density_correlation"]
    neg_density = f3["top_negative_density_correlation"]
    top_prog = f4["top_positive_pseudotime_programs"][0]
    early_prog = f4["top_negative_pseudotime_programs"][0]
    micro = f5["top_microenvironment_trend"]
    xenium = f7["top_xenium_trend"]

    return [
        {
            "type": "abstract",
            "title": "Abstract",
            "paragraphs": [
                (
                    "Spatial omics experiments now profile millions of cells in "
                    "intact tissues, but many analyses still separate cell "
                    "phenotype, morphology, extracellular matrix (ECM) structure, "
                    "and disease progression into disconnected notebooks. We "
                    "present SpatioEv, a reusable Python toolbox for spatial "
                    "evolution analysis in multiplexed imaging data with "
                    "extensions to spatial transcriptomics. SpatioEv organizes "
                    "segmentation quality control, phenotype feature engineering, "
                    "semi-supervised annotation refinement, density and "
                    "point-pattern statistics, local spatial autocorrelation, "
                    "niche-boundary analysis, graph summaries, cell-ECM coupling, "
                    "and pseudotime-linked microenvironment dynamics into an "
                    "auditable AnnData-centered workflow. We refactored the "
                    "codebase into a GitHub-ready package with lazy optional "
                    "dependencies, tests, a generated public function catalog, "
                    "and seven executable tutorials that use local example data "
                    "or synthetic fallbacks. Applying the current package to "
                    "pancreatic ductal adenocarcinoma (PDAC) multiplexed imaging "
                    f"data, SpatioEv processed an example field of {fmt_int(f2['exp2_total_cells'])} "
                    f"cells, quantified {fmt_int(f3['tiles'])} spatial tiles, and "
                    "identified an epithelial trajectory in which desmoplastic "
                    f"and invasive-niche features increased with pseudotime (Spearman r = {fmt_float(top_prog['spearman_r'])}) "
                    f"while early duct-like structure decreased (r = {fmt_float(early_prog['spearman_r'])}). "
                    "Across multiplexed and Xenium-derived analyses, fibroblast "
                    "abundance rose along progression, immune proximity changed "
                    "branch-specifically, and ECM-cell statistics recovered "
                    "disease-associated matrix-immune coupling in inflammatory "
                    "synovium. These examples show that SpatioEv is not only a "
                    "function collection but a reproducible framework for "
                    "turning high-dimensional tissue images into interpretable "
                    "models of spatial disease evolution."
                )
            ],
        },
        {
            "type": "section",
            "title": "Introduction",
            "paragraphs": [
                (
                    "Tissue biology is spatial by definition. Epithelial cells "
                    "remain constrained by ducts, glands, and basement membranes; "
                    "fibroblasts remodel matrix around lesions; immune cells are "
                    "recruited, excluded, or organized into local neighborhoods; "
                    "and disease progression often occurs at boundaries rather "
                    "than uniformly across a sample. Multiplexed imaging and "
                    "spatial transcriptomics preserve these relationships, but "
                    "their value depends on analysis frameworks that can move "
                    "from image-derived measurements to biological interpretation "
                    "without losing coordinate systems, cell identities, or "
                    "quality-control provenance."
                ),
                (
                    "Several influential platforms have made spatial omics more "
                    "tractable. Squidpy and Giotto provide broad spatial omics "
                    "analysis ecosystems; SpatialData establishes a standard "
                    "data framework for multimodal spatial datasets; and BANKSY "
                    "shows how local molecular neighborhoods can support scalable "
                    "cell typing and domain segmentation. These advances are "
                    "essential, but common disease projects still require custom "
                    "code for segmentation review, morphology-aware phenotype "
                    "refinement, multi-scale point-pattern statistics, ECM-cell "
                    "coupling, branch-aware microenvironment dynamics, and "
                    "publication-grade tables linking method output to biological "
                    "claims."
                ),
                (
                    "SpatioEv was developed from this practical gap. The toolbox "
                    "grew out of multiplexed imaging studies of PDAC, autoimmune "
                    "pancreatitis, and inflammatory synovium, where biological "
                    "questions often involved not only which cells are present, "
                    "but where they sit relative to ducts, stroma, matrix fibers, "
                    "immune aggregates, and pseudotime-like epithelial state "
                    "changes. In PDAC, dense stroma and cancer-associated "
                    "fibroblasts shape intratumoral architecture, immune "
                    "exclusion, and invasive epithelial programs. PanIN-to-PDAC "
                    "studies increasingly show that epithelial progression and "
                    "microenvironment remodeling are intertwined, branch-like, "
                    "and spatially heterogeneous rather than a single linear "
                    "clock. These are exactly the situations where explicit "
                    "spatial features and interpretable niche summaries are "
                    "needed."
                ),
                (
                    "Here we report the current SpatioEv package as a reusable "
                    "and reproducible toolbox. We rebuilt the repository for "
                    "GitHub upload, expanded tutorials into a guided user manual, "
                    "generated a public function catalog, and used the existing "
                    "example datasets to produce manuscript figures. We then "
                    "reinterpret the example data as a biological case study: "
                    "quality-controlled PDAC multiplexed imaging reveals "
                    "fibroblast-ductal organization, epithelial pseudotime is "
                    "accompanied by desmoplastic niche remodeling, ECM-cell "
                    "statistics capture matrix-associated inflammation, and "
                    "Xenium-derived transcriptomic outputs reproduce key "
                    "microenvironment trends."
                ),
            ],
        },
        {
            "type": "figure",
            "number": 1,
            "caption": (
                f"Figure 1. SpatioEv package architecture and reproducibility workflow. "
                f"The current package exposes {f1['public_functions_represented']} public "
                f"functions/classes across {f1['api_modules']} analysis modules, spanning "
                "quality control, phenotype modeling, spatial density/statistics, niche "
                "graphs, ECM-cell analysis, pseudotime dynamics, visualization, and "
                "Xenium-oriented extensions. The release workflow combines imports, unit "
                "tests, tutorial execution, figure generation, and package build checks."
            ),
        },
        {
            "type": "section",
            "title": "Results",
            "paragraphs": [],
        },
        {
            "type": "subsection",
            "title": "SpatioEv is organized as a reusable and testable spatial-evolution package",
            "paragraphs": [
                (
                    "The first goal was to convert a research codebase into a "
                    "package that users can install, inspect, and extend. The "
                    "source now lives under a conventional `spatioev/` package "
                    "tree with package metadata in `pyproject.toml`, an explicit "
                    "data policy, a release checklist, and tests that can run in "
                    "a lightweight checkout. Heavy tools such as Scanpy, scimap, "
                    "Napari, SpatialData, and Squidpy are optional extras rather "
                    "than unconditional import-time dependencies. This matters "
                    "for reproducibility: a user can import the core package, run "
                    "segmentation QC, and compute spatial statistics without "
                    "first resolving a full interactive-imaging stack."
                ),
                (
                    "The public API is intentionally grouped by analysis stage. "
                    "QC functions flag segmentation outliers; preprocessing "
                    "utilities move marker values and z-scored features into "
                    "AnnData observations; phenotype functions support SVM-based "
                    "classification and refinement; spatial functions compute "
                    "density, Ripley and Moran statistics, local co-localization, "
                    "niche boundaries, graph features, ECM-cell links, and "
                    "pseudotime dynamics. The tutorial generator writes a "
                    "function catalog so users can see where each function sits "
                    "in the biological workflow, instead of discovering the "
                    "package one import error at a time."
                ),
            ],
        },
        {
            "type": "subsection",
            "title": "Segmentation QC and phenotype maps establish an interpretable PDAC tissue scaffold",
            "paragraphs": [
                (
                    "Spatial interpretation begins with quality control because "
                    "segmentation artifacts create false neighbors, false "
                    "densities, and distorted morphology. Using the local PDAC "
                    f"example field, SpatioEv read {fmt_int(f2['exp2_total_cells'])} single-cell "
                    "objects and applied area and nuclear-to-cell ratio QC to a "
                    f"{fmt_int(f2['qc_sample_cells'])}-cell sample. The QC pass flagged "
                    f"{fmt_pct(f2['qc_percent_removed'])} of sampled cells under the example "
                    "thresholds, providing a transparent estimate of how much "
                    "segmentation uncertainty would otherwise enter downstream "
                    "statistics."
                ),
                (
                    f"The dominant phenotype in this example was {f2['top_phenotype']} "
                    f"({fmt_pct(100 * f2['top_phenotype_fraction'])} of cells), followed by "
                    "ductal epithelium, mesenchymal populations, endothelial "
                    "cells, acinar epithelium, and lymphoid lineages. This "
                    "composition is biologically consistent with PDAC tissue: "
                    "ductal epithelial structures are embedded in a large stromal "
                    "compartment, and the degree to which fibroblasts, endothelial "
                    "cells, and immune populations approach those ducts carries "
                    "information about tumor architecture and immune exclusion. "
                    "The point of the QC-to-phenotype stage is therefore not only "
                    "to clean a table, but to define the cellular scaffold on "
                    "which every spatial claim will rest."
                ),
            ],
        },
        {
            "type": "figure",
            "number": 2,
            "caption": (
                "Figure 2. Example segmentation quality control and phenotype "
                "architecture in PDAC multiplexed imaging. The panels summarize "
                "area and nuclear-to-cell ratio QC, phenotype composition, and a "
                "spatial phenotype map from the local exp_2 dataset. The example "
                f"QC thresholds removed {fmt_pct(f2['qc_percent_removed'])} of sampled "
                "objects, illustrating how SpatioEv surfaces segmentation risk "
                "before spatial statistics are computed."
            ),
        },
        {
            "type": "subsection",
            "title": "Multi-scale spatial statistics recover compartmentalized ductal-stromal organization",
            "paragraphs": [
                (
                    f"SpatioEv tiled the example image into {fmt_int(f3['tiles'])} spatial "
                    "bins and quantified phenotype-specific densities. The "
                    f"strongest positive density correlation was between {top_density['pair'][0]} "
                    f"and {top_density['pair'][1]} (r = {fmt_float(top_density['r'])}), while "
                    f"{neg_density['pair'][0]} and {neg_density['pair'][1]} were negatively "
                    f"correlated (r = {fmt_float(neg_density['r'])}). This pattern is "
                    "compatible with preserved or metaplastic epithelial regions "
                    "occupying different tissue space from fibroblast-rich "
                    "desmoplastic regions, rather than all epithelial populations "
                    "being uniformly mixed with stroma."
                ),
                (
                    "The cross-Ripley analysis adds a scale axis to the same "
                    "question. Ductal-fibroblast attraction peaked at an example "
                    f"radius of {fmt_int(f3['ripley_peak_radius_pixels'])} pixels, indicating "
                    "that fibroblast enrichment around ductal structures is not "
                    "only a global composition difference but an organized "
                    "neighborhood signal. In biological terms, this supports a "
                    "model in which epithelial compartments are surrounded by "
                    "stromal niches at characteristic spatial scales, the kind "
                    "of relationship expected in desmoplastic PDAC and PanIN-like "
                    "remodeling."
                ),
            ],
        },
        {
            "type": "figure",
            "number": 3,
            "caption": (
                "Figure 3. Density and spatial statistics quantify tissue "
                "organization beyond visual inspection. Tile-density maps, "
                "phenotype density correlations, and cross-Ripley curves show "
                "where epithelial and stromal phenotypes co-occupy or separate "
                "in physical space. The strongest positive density correlation "
                f"was {top_density['pair'][0]} versus {top_density['pair'][1]} "
                f"(r = {fmt_float(top_density['r'])}); the strongest negative correlation "
                f"was {neg_density['pair'][0]} versus {neg_density['pair'][1]} "
                f"(r = {fmt_float(neg_density['r'])})."
            ),
        },
        {
            "type": "subsection",
            "title": "Niche-level pseudotime links epithelial state to desmoplastic progression",
            "paragraphs": [
                (
                    f"Across the combined multiplexed imaging example, SpatioEv summarized "
                    f"{fmt_int(f4['n_niches'])} epithelial niches and retained {f4['n_pseudotime_branches']} "
                    "branch annotations. The strongest positive pseudotime "
                    f"program was {top_prog['label']} (Spearman r = {fmt_float(top_prog['spearman_r'])}), "
                    f"whereas {early_prog['label']} decreased along the same axis "
                    f"(r = {fmt_float(early_prog['spearman_r'])}). The direction of these "
                    "signals is biologically plausible: as ductal structures "
                    "move away from an early duct-like state, their surrounding "
                    "niches acquire features associated with invasion, "
                    "desmoplasia, and higher-grade PanIN/PDAC-like organization."
                ),
                (
                    "This interpretation must be stated carefully. SpatioEv "
                    "pseudotime is not a literal chronological clock and should "
                    "not imply that every PanIN-like niche inevitably becomes "
                    "invasive cancer. The literature increasingly supports a "
                    "branched model in which epithelial state, ductal geometry, "
                    "stromal reaction, and immune context evolve asynchronously. "
                    "The package therefore treats pseudotime as an interpretable "
                    "state manifold: it orders niches by morphology, epithelial "
                    "program, and microenvironmental context, then lets users "
                    "test whether biological features change globally, "
                    "branch-specifically, or abruptly across transitions."
                ),
            ],
        },
        {
            "type": "figure",
            "number": 4,
            "caption": (
                "Figure 4. Niche-level pseudotime analysis identifies spatial "
                "programs associated with PDAC-like progression. The example "
                f"analysis summarized {fmt_int(f4['n_niches'])} niches across {f4['n_pseudotime_branches']} "
                "branches. Invasion/desmoplasia increased with pseudotime, "
                "while early duct-like organization decreased, supporting a "
                "branched epithelial-state interpretation rather than a purely "
                "linear progression model."
            ),
        },
        {
            "type": "subsection",
            "title": "Microenvironment dynamics identify branch-specific fibroblast and immune remodeling",
            "paragraphs": [
                (
                    "The most prominent multiplexed microenvironment trend was "
                    f"{micro['label'].lower()}, which increased with pseudotime "
                    f"(Spearman r = {fmt_float(micro['spearman_r'])}; late-minus-early median = "
                    f"{fmt_float(micro['late_minus_early_median'])}). This trend "
                    "matches a central feature of PDAC biology: epithelial "
                    "progression is accompanied by increasing stromal investment "
                    "and fibroblast-rich desmoplasia. Because SpatioEv stores "
                    "these trends as tables, users can examine not only global "
                    "changes but the specific phenotypes, markers, and interaction "
                    "features driving each trajectory segment."
                ),
                (
                    f"Branch-time analysis found {f5['n_transition_events']} transition events. "
                    f"The strongest event was {f5['strongest_transition']['transition']} in the "
                    f"{f5['strongest_transition']['dataset']} analysis, where the top "
                    f"changes included: {f5['strongest_transition']['top_changes']}. "
                    "This is the type of biological signal that motivates "
                    "SpatioEv: rather than reporting that a sample is simply "
                    "immune-rich or immune-poor, the workflow localizes when "
                    "immune proximity changes relative to epithelial progression "
                    "and which branches carry that remodeling."
                ),
            ],
        },
        {
            "type": "figure",
            "number": 5,
            "caption": (
                "Figure 5. SpatioEv maps microenvironment dynamics along "
                "epithelial pseudotime. Fibroblast proportion increased most "
                f"strongly with pseudotime (r = {fmt_float(micro['spearman_r'])}), while "
                "branch-time event analysis highlighted sharp local changes in "
                "ductal-immune proximity. These outputs turn tissue composition "
                "into branch-aware hypotheses about stromal activation and immune "
                "redistribution."
            ),
        },
        {
            "type": "subsection",
            "title": "ECM-cell statistics extend SpatioEv from cell neighborhoods to matrix biology",
            "paragraphs": [
                (
                    "Cell neighborhoods alone cannot capture the physical "
                    "barriers and migration tracks created by the extracellular "
                    "matrix. SpatioEv therefore includes functions to build "
                    "cell-fiber links, estimate nearest cell-to-fiber distances, "
                    "count fiber density near cells, compute cell-ECM cross-Ripley "
                    "and cross-Moran statistics, map cell-derived signals onto "
                    "fibers, and identify ECM-cell graph niches. This makes ECM "
                    "a first-class spatial object rather than a background image "
                    "layer."
                ),
                (
                    "The inflammatory synovium example illustrates why this "
                    "matters beyond PDAC. In the RA/OA analysis, the largest "
                    f"density shift was {f6['top_density_shift']['phenotype']} near "
                    f"{f6['top_density_shift']['fiber_type']} fibers (RA-minus-OA = "
                    f"{fmt_float(f6['top_density_shift']['ra_minus_oa'], 6)}). The largest "
                    f"nearest-distance shift was {f6['top_distance_shift']['phenotype']} relative "
                    f"to {f6['top_distance_shift']['fiber_type']} (RA-minus-OA = "
                    f"{fmt_float(f6['top_distance_shift']['ra_minus_oa'])}), and the strongest "
                    f"alignment coupling involved {f6['top_alignment_coupling']['interaction']} "
                    f"(cross-Moran's I = {fmt_float(f6['top_alignment_coupling']['cross_morans_i'])}). "
                    "These patterns are consistent with the view that inflammatory "
                    "tissue organization is jointly encoded by immune-cell "
                    "position and matrix architecture, including fibronectin- and "
                    "collagen-associated microenvironments."
                ),
            ],
        },
        {
            "type": "figure",
            "number": 6,
            "caption": (
                "Figure 6. ECM-cell analysis links fiber architecture to immune "
                "and stromal organization. The example RA/OA outputs compare "
                "cell density near fiber types, nearest cell-to-fiber distances, "
                "pathology-specific interaction matrices, and local ECM-cell "
                "coupling. These panels demonstrate how the same SpatioEv logic "
                "can analyze non-cancer tissue remodeling without changing the "
                "underlying package API."
            ),
        },
        {
            "type": "subsection",
            "title": "Xenium outputs reproduce key progression trends in a transcriptomic spatial modality",
            "paragraphs": [
                (
                    "Although SpatioEv was built around multiplexed imaging, the "
                    "current repository includes Xenium-oriented utilities for "
                    "data audit, DAPI morphology extraction, cell annotation, "
                    "niche features, BANKSY integration, and pseudotime-linked "
                    "interaction summaries. The example Xenium analysis covered "
                    f"{f7['xenium_samples']} samples and {f7['xenium_branches']} branches. "
                    f"Fibroblast proportion again increased with pseudotime "
                    f"(Spearman r = {fmt_float(xenium['spearman_r'])}; late-minus-early median = "
                    f"{fmt_float(xenium['late_minus_early_median'])}), providing cross-modality "
                    "support for the fibroblast/desmoplasia signal observed in "
                    "multiplexed imaging."
                ),
                (
                    "The branch labels are biologically interpretable rather "
                    "than arbitrary clusters. Example branches include residual "
                    "normal duct-like states, immune-excluded gland-poor or "
                    "undifferentiated-like tumor contexts, and more glandular or "
                    "differentiated tumor architecture. Ligand-receptor and "
                    "neighborhood-program summaries further suggest that "
                    "epithelial, fibroblast, endothelial, and immune features "
                    "change with branch-time. Because the Xenium panel is "
                    "targeted, these signals should be interpreted as proxy "
                    "programs rather than exhaustive transcriptomic states, but "
                    "they show that SpatioEv's niche-pseudotime concepts can "
                    "transfer across modality boundaries."
                ),
            ],
        },
        {
            "type": "figure",
            "number": 7,
            "caption": (
                "Figure 7. Xenium extension of SpatioEv spatial-evolution logic. "
                f"The example analysis contains {f7['xenium_samples']} samples and "
                f"{f7['xenium_branches']} branches, with fibroblast proportion increasing "
                f"along pseudotime (r = {fmt_float(xenium['spearman_r'])}). BANKSY-derived "
                "domain summaries and branch biology annotations provide a "
                "transcriptomic counterpart to the multiplexed imaging trajectory."
            ),
        },
        {
            "type": "section",
            "title": "Discussion",
            "paragraphs": [
                (
                    "SpatioEv is designed for the stage of spatial biology where "
                    "the question is no longer only which cell types are present, "
                    "but how tissue architecture changes. The current package "
                    "turns segmentation QC, phenotyping, density maps, local "
                    "statistics, niche boundaries, cell graphs, ECM-cell links, "
                    "and pseudotime dynamics into a coherent workflow. The "
                    "example analyses show that this workflow can recover "
                    "biologically credible patterns: fibroblast-rich stroma "
                    "increases with PDAC-like epithelial progression, ductal and "
                    "acinar compartments separate from fibroblast-dense regions, "
                    "immune proximity changes at branch-specific transitions, "
                    "and ECM architecture contributes measurable information in "
                    "inflammatory tissue."
                ),
                (
                    "The main conceptual contribution is interpretability. Many "
                    "spatial workflows produce clusters or embeddings that are "
                    "useful but difficult to audit. SpatioEv emphasizes "
                    "intermediate tables: QC summaries, phenotype probabilities, "
                    "density matrices, Ripley curves, Moran scores, local "
                    "neighbor summaries, graph features, ECM-cell link tables, "
                    "and branch-time event tables. These tables make it possible "
                    "to move from a figure to a testable biological statement, "
                    "and then back to the cells, fibers, or niches supporting "
                    "that statement."
                ),
                (
                    "The current release also has clear limitations. The example "
                    "PDAC analyses are demonstration datasets rather than a "
                    "validated clinical cohort. Pseudotime should be treated as "
                    "a spatial state ordering, not a direct chronological lineage. "
                    "Xenium panels are targeted and cannot support every canonical "
                    "PDAC subtype or PanIN module. The WGCNA/WFCNA-like feature "
                    "network workflow described in earlier manuscript drafts is "
                    "not included as a current public API because it has not yet "
                    "been rebuilt, tested, and documented in the package. It is "
                    "best framed as a future extension that could turn the "
                    "existing niche feature tables into co-varying spatial "
                    "feature modules."
                ),
                (
                    "Future work should add explicit duct/lumen geometry, "
                    "epithelial boundary roughness, periductal matrix orientation, "
                    "and H&E-derived stromal texture features, all of which are "
                    "strongly motivated by PanIN and PDAC spatial pathology. "
                    "Continuous integration on redistributable public datasets "
                    "will also be important before a stable release. Nevertheless, "
                    "the current package already provides a practical, tested, "
                    "and biologically motivated toolbox for researchers who need "
                    "to analyze spatial evolution in multiplexed imaging and "
                    "spatial transcriptomics."
                ),
            ],
        },
        {
            "type": "section",
            "title": "Methods",
            "paragraphs": [
                (
                    "Package organization. SpatioEv is a Python package organized "
                    "into modules for I/O, configuration, segmentation QC, "
                    "preprocessing, machine-learning phenotype models, annotation "
                    "refinement, visualization, spatial density, spatial "
                    "statistics, niche boundaries, cell graphs, ECM-cell analysis, "
                    "pixel-feature extraction, and pseudotime dynamics. Optional "
                    "dependencies are grouped in `pyproject.toml` and imported "
                    "lazily where possible."
                ),
                (
                    "Tutorial generation. The tutorial series was generated by "
                    "`scripts/write_tutorial_notebooks.py`, which also writes "
                    "`docs/function_catalog.csv` and `docs/function_catalog.md`. "
                    "The notebooks use the existing local example H5AD and CSV "
                    "annotations when present and otherwise construct a small "
                    "synthetic AnnData object so that the tutorials remain "
                    "runnable in GitHub clones without private or large data."
                ),
                (
                    "QC and phenotype analysis. Segmentation QC used cell area, "
                    "physical area conversion, and nuclear-to-cell ratio columns "
                    "stored in `adata.obs`. The manuscript example used pixel "
                    "size 0.325, minimum area 5, maximum area 650, and maximum "
                    "nuclear-to-cell ratio 1.0 on a sampled exp_2 table. Phenotype "
                    "composition was computed from the Tier_A annotation column "
                    "after alignment with the local annotation CSV."
                ),
                (
                    "Spatial density and statistics. Tile-level density used a "
                    "1024-pixel grid over the exp_2 example field. Phenotype "
                    "density correlations were computed across tile-by-phenotype "
                    "density matrices. Cross-Ripley curves sampled source and "
                    "target cells with fixed random seeds and evaluated ductal-"
                    "fibroblast organization across increasing radii."
                ),
                (
                    "Niche and pseudotime analysis. Combined epithelial-niche "
                    "tables from exp_2 through exp_5 provided pooled pseudotime, "
                    "branch labels, and pathology module scores. Spearman "
                    "correlations related module scores and microenvironment "
                    "features to pseudotime. Branch-time event tables identified "
                    "features with the largest standardized changes between "
                    "successive pseudotime bins within each branch."
                ),
                (
                    "ECM-cell analysis. ECM examples used existing RA/OA result "
                    "tables under `paper/notebooks/results/ra_oa_ecm_cell/`. The package "
                    "API supports construction of cell-fiber links, nearest "
                    "cell-to-fiber distances, fiber density near cells, cell-ECM "
                    "cross-Ripley and cross-Moran statistics, cell-to-fiber signal "
                    "mapping, ECM graph niches, and radius-based ECM-cell "
                    "neighborhood clustering."
                ),
                (
                    "Xenium extension. Xenium examples used pseudotime and branch "
                    "summary tables under `data/xenium_pancreas_10x/pseudotime/` "
                    "and BANKSY integration outputs under the notebook results "
                    "directory. Targeted ligand-receptor and neighborhood-program "
                    "signals were interpreted as panel-limited proxies."
                ),
                (
                    "Figure generation and verification. Manuscript figures were "
                    "generated by `scripts/generate_manuscript_figures.py`, which "
                    "writes PNG panels, analysis tables, and `analysis_summary.json`. "
                    "Tutorial notebooks were executed with nbconvert as a "
                    "smoke-test of the documented workflows."
                ),
            ],
        },
        {
            "type": "section",
            "title": "Data and code availability",
            "paragraphs": [
                (
                    "The package source, tests, tutorials, release checklist, "
                    "and manuscript-generation scripts are contained in the "
                    "SpatioEv repository. Large local data directories, H5AD "
                    "files, Zarr stores, raw images, and generated result folders "
                    "are intentionally excluded from Git history by the data "
                    "policy. Redistributable example data should be released "
                    "through an approved archive or GitHub release asset with "
                    "checksums before public upload."
                )
            ],
        },
        {
            "type": "section",
            "title": "References",
            "paragraphs": REFERENCES,
        },
    ]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "D9E2EC")


def configure_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.10


def add_title_block(doc: Document) -> None:
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(6)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run(TITLE)
    run.font.name = "Calibri"
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string("0B2545")

    author = doc.add_paragraph()
    author.paragraph_format.space_after = Pt(3)
    run = author.add_run(AUTHORS)
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.bold = True

    aff = doc.add_paragraph()
    aff.paragraph_format.space_after = Pt(10)
    run = aff.add_run(AFFILIATION)
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string("555555")

    keys = doc.add_paragraph()
    keys.paragraph_format.space_after = Pt(10)
    run = keys.add_run(f"Keywords: {KEYWORDS}")
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    run.font.italic = True


def add_key_results_table(doc: Document, summary: dict) -> None:
    rows = [
        ("Example imaging cells", fmt_int(summary["figure_2"]["exp2_total_cells"])),
        ("QC objects flagged in sample", fmt_pct(summary["figure_2"]["qc_percent_removed"])),
        ("Spatial density tiles", fmt_int(summary["figure_3"]["tiles"])),
        (
            "Pseudotime niches",
            f"{fmt_int(summary['figure_4']['n_niches'])} niches across {summary['figure_4']['n_pseudotime_branches']} branches",
        ),
        (
            "Top progression signal",
            f"{summary['figure_4']['top_positive_pseudotime_programs'][0]['label']} (r = {fmt_float(summary['figure_4']['top_positive_pseudotime_programs'][0]['spearman_r'])})",
        ),
        (
            "Top microenvironment signal",
            f"{summary['figure_5']['top_microenvironment_trend']['label']} (r = {fmt_float(summary['figure_5']['top_microenvironment_trend']['spearman_r'])})",
        ),
        (
            "Xenium cross-modality signal",
            f"{summary['figure_7']['top_xenium_trend']['label']} (r = {fmt_float(summary['figure_7']['top_xenium_trend']['spearman_r'])})",
        ),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.columns[0].width = Inches(2.4)
    table.columns[1].width = Inches(4.0)
    set_table_borders(table)

    hdr = table.rows[0].cells
    hdr[0].text = "Claim"
    hdr[1].text = "Current example-data support"
    for cell in hdr:
        set_cell_shading(cell, "F2F4F7")
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    for claim, value in rows:
        cells = table.add_row().cells
        cells[0].text = claim
        cells[1].text = value
        for cell in cells:
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(4)


def add_figure_docx(doc: Document, number: int, caption: str) -> None:
    path = figure_path(number)
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run()
    run.add_picture(str(path), width=Inches(6.45))

    cap = doc.add_paragraph()
    cap.paragraph_format.space_after = Pt(10)
    cap.paragraph_format.keep_with_next = False
    run = cap.add_run(caption)
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor.from_string("333333")


def write_docx(content: list[dict], summary: dict) -> None:
    doc = Document()
    configure_doc(doc)
    add_title_block(doc)
    add_key_results_table(doc, summary)

    for block in content:
        block_type = block["type"]
        if block_type == "abstract":
            doc.add_heading(block["title"], level=1)
            for text in block["paragraphs"]:
                doc.add_paragraph(text)
        elif block_type == "section":
            doc.add_heading(block["title"], level=1)
            for text in block["paragraphs"]:
                doc.add_paragraph(text)
        elif block_type == "subsection":
            doc.add_heading(block["title"], level=2)
            for text in block["paragraphs"]:
                doc.add_paragraph(text)
        elif block_type == "figure":
            add_figure_docx(doc, block["number"], block["caption"])
        else:
            raise ValueError(f"Unknown block type: {block_type}")

    DOCX_OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(DOCX_OUT)


def write_markdown(content: list[dict], summary: dict) -> None:
    lines = [
        f"# {TITLE}",
        "",
        AUTHORS,
        "",
        AFFILIATION,
        "",
        f"**Keywords:** {KEYWORDS}",
        "",
        "## Key Results From Current Example Data",
        "",
        "| Claim | Current example-data support |",
        "|---|---|",
        f"| Example imaging cells | {fmt_int(summary['figure_2']['exp2_total_cells'])} |",
        f"| QC objects flagged in sample | {fmt_pct(summary['figure_2']['qc_percent_removed'])} |",
        f"| Spatial density tiles | {fmt_int(summary['figure_3']['tiles'])} |",
        f"| Pseudotime niches | {fmt_int(summary['figure_4']['n_niches'])} niches across {summary['figure_4']['n_pseudotime_branches']} branches |",
        f"| Top progression signal | {summary['figure_4']['top_positive_pseudotime_programs'][0]['label']} (r = {fmt_float(summary['figure_4']['top_positive_pseudotime_programs'][0]['spearman_r'])}) |",
        f"| Top microenvironment signal | {summary['figure_5']['top_microenvironment_trend']['label']} (r = {fmt_float(summary['figure_5']['top_microenvironment_trend']['spearman_r'])}) |",
        f"| Xenium cross-modality signal | {summary['figure_7']['top_xenium_trend']['label']} (r = {fmt_float(summary['figure_7']['top_xenium_trend']['spearman_r'])}) |",
        "",
    ]

    for block in content:
        block_type = block["type"]
        if block_type in {"abstract", "section"}:
            lines.append(f"## {block['title']}")
            lines.append("")
            for text in block["paragraphs"]:
                lines.append(dedent(text).strip())
                lines.append("")
        elif block_type == "subsection":
            lines.append(f"### {block['title']}")
            lines.append("")
            for text in block["paragraphs"]:
                lines.append(dedent(text).strip())
                lines.append("")
        elif block_type == "figure":
            rel = figure_path(block["number"]).relative_to(OUT_DIR).as_posix()
            lines.append(f"![Figure {block['number']}]({rel})")
            lines.append("")
            lines.append(f"*{block['caption']}*")
            lines.append("")
        else:
            raise ValueError(f"Unknown block type: {block_type}")

    MD_OUT.write_text("\n".join(lines).strip() + "\n", encoding="utf-8")


def main() -> None:
    summary = load_summary()
    content = build_content(summary)
    write_markdown(content, summary)
    write_docx(content, summary)
    print(f"Wrote {MD_OUT}")
    print(f"Wrote {DOCX_OUT}")


if __name__ == "__main__":
    main()
